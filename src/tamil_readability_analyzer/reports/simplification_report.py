"""Word document export for simplification reports."""

import io
import re


def generate_simplification_docx(report):
    """Build the simplification DOCX and return bytes plus response headers."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    target_grade = report.get('target_grade', 3)
    sentences = report.get('sentences', [])

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    title = doc.add_heading(level=1)
    title.clear()
    run = title.add_run(f'Tamil Text Simplification Report - Target: Standard {target_grade}')
    run.font.size = Pt(14)

    summary = doc.add_paragraph()
    summary.add_run(
        f'Total sentences: {report.get("total_sentences", 0)}  |  '
        f'Total Tamil words: {report.get("total_words", 0)}  |  '
        f'Hard words: {report.get("hard_word_count", 0)} '
        f'({report.get("hard_word_pct", 0):.1f}%)  |  '
        f'Complex sentences: {report.get("complex_sentence_count", 0)}'
    ).font.size = Pt(10)
    doc.add_paragraph()

    final_text = (report.get('final_text') or '').strip()
    if final_text:
        doc.add_heading('Edited simplified draft', level=2)
        for para in final_text.splitlines():
            if para.strip():
                p = doc.add_paragraph()
                r = p.add_run(para.strip())
                r.font.size = Pt(11)
        doc.add_page_break()

    red = RGBColor(0xC0, 0x39, 0x2B)
    amber = RGBColor(0xB8, 0x7A, 0x12)
    green = RGBColor(0x0F, 0x6E, 0x56)
    grey = RGBColor(0x88, 0x88, 0x88)
    blue = RGBColor(0x1A, 0x5F, 0xA8)

    for index, sent in enumerate(sentences, 1):
        if not sent.get('is_complex'):
            p = doc.add_paragraph()
            r = p.add_run(sent['original'])
            r.font.size = Pt(11)
            r.font.color.rgb = grey
            continue

        doc.add_heading(f'Sentence {index}', level=3)

        orig_para = doc.add_paragraph()
        orig_para.paragraph_format.space_after = Pt(2)
        original = sent['original']
        hard_set = {h['word'] for h in sent.get('hard_words', [])}

        pos = 0
        for match in re.finditer(r'[\u0B80-\u0BFF]{2,}', original):
            if match.start() > pos:
                r = orig_para.add_run(original[pos:match.start()])
                r.font.size = Pt(11)
            word = match.group()
            r = orig_para.add_run(word)
            r.font.size = Pt(11)
            if word in hard_set:
                r.font.color.rgb = red
                r.font.bold = True
            pos = match.end()
        if pos < len(original):
            r = orig_para.add_run(original[pos:])
            r.font.size = Pt(11)

        if sent.get('hard_words'):
            doc.add_paragraph('Word suggestions:', style='List Bullet')
            for hard_word in sent['hard_words']:
                sug_para = doc.add_paragraph(style='List Bullet 2')
                r = sug_para.add_run(hard_word['word'])
                r.font.color.rgb = red
                r.font.bold = True
                r.font.size = Pt(10)
                grade_label = f" (Std {hard_word['grade']})" if hard_word['grade'] else " (unknown grade)"
                sug_para.add_run(grade_label).font.size = Pt(9)
                sug_para.add_run('  ->  ').font.size = Pt(10)
                if hard_word['suggestions']:
                    best = hard_word['suggestions'][0]
                    rb = sug_para.add_run(best['stem'])
                    rb.font.color.rgb = green
                    rb.font.bold = True
                    rb.font.size = Pt(10)
                    rb2 = sug_para.add_run(f" (Std {best['grade']})")
                    rb2.font.size = Pt(9)
                    for alt in hard_word['suggestions'][1:]:
                        sug_para.add_run(',  ').font.size = Pt(10)
                        ra = sug_para.add_run(alt['stem'])
                        ra.font.color.rgb = amber
                        ra.font.size = Pt(10)
                        ra2 = sug_para.add_run(f" (Std {alt['grade']})")
                        ra2.font.size = Pt(9)
                else:
                    r_na = sug_para.add_run('No suggestion found - manual rewrite needed')
                    r_na.font.color.rgb = amber
                    r_na.font.size = Pt(9)

        if sent.get('rewritten') and sent['rewritten'] != sent['original']:
            rw_para = doc.add_paragraph()
            rw_para.add_run('Suggested rewrite:  ').font.size = Pt(9)
            rr = rw_para.add_run(sent['rewritten'])
            rr.font.size = Pt(11)
            rr.font.color.rgb = green

        if sent.get('split_suggestion'):
            sp_para = doc.add_paragraph()
            sp_para.add_run('Split into shorter sentences:  ').font.size = Pt(9)
            for part in sent['split_suggestion']:
                rp = sp_para.add_run(part + '  ')
                rp.font.size = Pt(11)
                rp.font.color.rgb = blue

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), {
        'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'Content-Disposition': f'attachment; filename="simplification_std{target_grade}.docx"',
    }
