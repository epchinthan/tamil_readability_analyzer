import os, re, json, sqlite3, datetime, io, math, hashlib, threading, logging, uuid, time
import werkzeug.utils
import analytics as _analytics
import meaning_kb as _meaning_kb
import suitability_engine as _suitability
import v27_offline as _v27
import textbook_importer as _importer
import reading_asr as _reading_asr
import reading_score as _reading_score
import tamil_features as _tamil_features

# Suppress pdfminer's "Cannot set [non-]stroke color" warnings.
# These appear when parsing PDFs that use DeviceN/Spot colorspaces (common in
# professionally printed textbooks). They are harmless — text extracts correctly.
logging.getLogger('pdfminer').setLevel(logging.ERROR)
from flask import Flask, request, jsonify, render_template, send_file
from ollama_client import ollama_health, tamil_author_rewrite, tamil_simple_explanation, tamil_lesson_plan, tamil_questions, generate, DEFAULT_MODEL, DEFAULT_BASE_URL
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024
app.config['UPLOAD_FOLDER'] = 'uploads'
DB_PATH = 'tamil_analyzer.db'
WATCH_FOLDER = None   # set via /api/config or config.json at startup

# ── Tamil Morphological Stemmer ───────────────────────────────────────────────

try:
    import snowballstemmer as _sb
    _SNOWBALL = _sb.stemmer('tamil')
    SNOWBALL_AVAILABLE = True
except ImportError:
    SNOWBALL_AVAILABLE = False

TAMIL_SUFFIXES = [
    'க்காகவே','யிலிருந்து','இலிருந்து','லிருந்து','க்கிடையே',
    'யிடமிருந்து','இடமிருந்து','கிடையே','நடுவே',
    'உடைய','க்காக','யிடம்','இடம்','உடன்','யோடு','ஓடு',
    'யினால்','இனால்','ஆலே','னால்','ஆல்',
    'யில்','இல்','ல்',
    'யிற்கே','யிற்கு','இற்கு','க்கு',
    'யின்','தின்','ரின்','இன்',
    'ரது','தது','அது','கள்','இருந்து','யே','வே','ஏ',
]
SANDHI_MAP = [
    (r'ட்டி$','டு'),(r'ட்டு$','டு'),(r'ட்ட$','டு'),
    (r'ண்ண$','ண்'),(r'ல்ல$','ல்'),(r'ன்ன$','ன்'),
    (r'ற்றி$','று'),(r'ற்ற$','று'),
]

def _strip_suffixes(word):
    changed = True
    while changed:
        changed = False
        for s in TAMIL_SUFFIXES:
            if word.endswith(s) and len(word)-len(s) >= 2:
                word = word[:-len(s)]; changed = True; break
    return word

def _apply_sandhi(word):
    for p, r in SANDHI_MAP:
        word = re.sub(p, r, word)
    return word

def stem_tamil_word(word):
    """Safely stem a Tamil token.

    Some PDFs/OCR outputs contain malformed Tamil Unicode sequences, stray
    combining marks, or zero-length tokens after cleanup. The Snowball Tamil
    stemmer can raise `string index out of range` on those edge cases.
    Never let one bad token fail the whole textbook.
    """
    word = (word or '').strip()
    if len(word) < 2 or not re.search(r'[\u0B80-\u0BFF]', word):
        return word

    try:
        sb = _SNOWBALL.stemWord(word) if SNOWBALL_AVAILABLE else word
    except Exception:
        sb = word

    if sb and sb != word:
        return _apply_sandhi(sb)

    stripped = _strip_suffixes(word)
    if stripped != word:
        if SNOWBALL_AVAILABLE:
            try:
                stripped = _SNOWBALL.stemWord(stripped)
            except Exception:
                pass
        return _apply_sandhi(stripped or word)
    return word

_stem_cache = {}
def get_stem(w):
    w = (w or '').strip()
    if not w:
        return ''
    if w not in _stem_cache:
        try:
            _stem_cache[w] = stem_tamil_word(w)
        except Exception as e:
            logging.getLogger('app').warning(f'Stemmer skipped malformed token {w!r}: {e}')
            _stem_cache[w] = w
    return _stem_cache[w]

# ── Sentence Analysis ─────────────────────────────────────────────────────────

def tokenize_sentences(text):
    """Split text into sentences on Tamil/Latin punctuation."""
    parts = re.split(r'[.!?।\u0964\u0965\n]+', text)
    return [p.strip() for p in parts if p.strip()]

def sentence_word_counts(text):
    """Return list of Tamil word counts per sentence."""
    counts = []
    for sent in tokenize_sentences(text):
        n = len(re.findall(r'[\u0B80-\u0BFF]{2,}', sent))
        if n > 0:
            counts.append(n)
    return counts

def sentence_stats(counts):
    """Compute max, avg, median from a list of sentence word counts."""
    if not counts:
        return {'max': 0, 'avg': 0.0, 'median': 0, 'total_sentences': 0}
    return {
        'max': max(counts),
        'avg': round(sum(counts) / len(counts), 1),
        'median': sorted(counts)[len(counts)//2],
        'total_sentences': len(counts),
    }

# ── Proper Noun Detection ─────────────────────────────────────────────────────

PLACE_SUFFIXES = [
    'நகர்','நகரம்','பட்டணம்','பட்டினம்','பூர்','பூரம்',
    'ஆறு','நதி','மலை','குன்று','குன்றம்','மேடு',
    'வாயில்','துறை','கரை','பாளையம்','பாடி',
    'கோட்டை','நல்லூர்','ஊர்','கிராமம்',
    'தீவு','கடல்','வனம்','காடு','சாலை',
    'மாவட்டம்','மாநிலம்','நாடு','தலைநகர்',
]
PERSON_NAME_SUFFIXES = [
    'ராஜ்','ராஜா','ராணி','குமார்','குமாரன்','குமாரி',
    'சேகர்','மோகன்','பாபு','தேவி','வேணி','லட்சுமி',
    'முருகன்','கணேஷ்','ராமன்','கிருஷ்ணன்','வேலன்',
    'அம்மன்','அம்மாள்','பிள்ளை','ஐயர்','ஐயங்கார்',
    'ஆச்சார்','நாடார்','தேவர்','மூர்த்தி',
]
DEITY_NAMES = {
    'முருகன்','கணேஷ்','கணபதி','விநாயகர்','ஆண்டவன்',
    'அல்லா','இயேசு','சிவன்','விஷ்ணு','பிரம்மன்',
    'சரஸ்வதி','லட்சுமி','பார்வதி','துர்கா','காளி',
    'ராமன்','கிருஷ்ணன்','அய்யப்பன்','கார்த்திகேயன்',
    'வேலன்','அம்மன்','மாரியம்மன்','காவேரி',
    'புத்தர்','மகாவீரர்','நாராயணன்',
}
FOREIGN_CLUSTERS = ['க்ஸ்','ட்ர','ப்ர','ஸ்ட்','ல்ட்','ன்ஸ்','ஸ்','ஜ்','ஹ்','ஃ']
LOAN_STARTERS    = ['ஸ','ஜ','ஹ','ஃ','ஷ','க்ஷ']

def detect_proper_nouns(stem_freq, grade_vocab_union):
    flagged = {}
    for word, freq in stem_freq.items():
        if word in grade_vocab_union: continue
        reasons = []
        if word in DEITY_NAMES: reasons.append('deity/god name')
        if any(word.endswith(s) for s in PLACE_SUFFIXES): reasons.append('place-name suffix')
        if any(word.endswith(s) for s in PERSON_NAME_SUFFIXES): reasons.append('person-name suffix')
        if any(c in word for c in FOREIGN_CLUSTERS): reasons.append('foreign name pattern')
        if any(word.startswith(s) for s in LOAN_STARTERS): reasons.append('loan-word starter')
        if freq <= 3 and not reasons: reasons.append('rare unknown word')
        if reasons: flagged[word] = reasons
    return flagged

# ── Database ──────────────────────────────────────────────────────────────────

def get_db():
    conn = sqlite3.connect(DB_PATH, timeout=30)
    conn.row_factory = sqlite3.Row
    # Performance PRAGMAs applied to every connection
    conn.execute('PRAGMA journal_mode=WAL')       # concurrent reads during writes
    conn.execute('PRAGMA synchronous=NORMAL')     # safe but faster than FULL
    conn.execute('PRAGMA cache_size=-64000')      # 64 MB page cache
    conn.execute('PRAGMA temp_store=MEMORY')      # temp tables in RAM
    conn.execute('PRAGMA mmap_size=268435456')    # 256 MB memory-mapped I/O
    return conn

def init_db():
    conn = get_db()
    conn.executescript('''
        CREATE TABLE IF NOT EXISTS grade_words (
            grade INTEGER NOT NULL, word TEXT NOT NULL,
            PRIMARY KEY (grade, word)
        );
        -- One row per file (multiple files per grade supported)
        CREATE TABLE IF NOT EXISTS grade_files (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            grade        INTEGER NOT NULL,
            filepath     TEXT NOT NULL UNIQUE,
            filename     TEXT NOT NULL,
            file_hash    TEXT NOT NULL,
            word_count   INTEGER DEFAULT 0,
            raw_count    INTEGER DEFAULT 0,
            sent_max     INTEGER DEFAULT 0,
            sent_avg     REAL    DEFAULT 0,
            sent_median  INTEGER DEFAULT 0,
            sent_total   INTEGER DEFAULT 0,
            source       TEXT    DEFAULT 'folder',
            processed_at TEXT
        );
        -- Aggregated view per grade (computed from grade_files)
        CREATE TABLE IF NOT EXISTS grade_meta (
            grade        INTEGER PRIMARY KEY,
            file_count   INTEGER DEFAULT 0,
            word_count   INTEGER DEFAULT 0,
            raw_count    INTEGER DEFAULT 0,
            sent_max     INTEGER DEFAULT 0,
            sent_avg     REAL    DEFAULT 0,
            sent_median  INTEGER DEFAULT 0,
            sent_total   INTEGER DEFAULT 0,
            updated_at   TEXT
        );
        CREATE TABLE IF NOT EXISTS word_grade_map (
            stem         TEXT PRIMARY KEY,
            first_grade  INTEGER NOT NULL
        );
        CREATE TABLE IF NOT EXISTS analyses (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            book_name       TEXT,
            analyzed_at     TEXT,
            total_words     INTEGER,
            unique_words    INTEGER,
            unique_stems    INTEGER,
            proper_nouns    TEXT,
            sentence_json   TEXT,
            results_json    TEXT
        );
        CREATE TABLE IF NOT EXISTS pending_extractions (
            id               INTEGER PRIMARY KEY AUTOINCREMENT,
            book_name        TEXT,
            created_at       TEXT,
            total_words      INTEGER,
            unique_words     INTEGER,
            unique_stems     INTEGER,
            stem_to_original TEXT,
            stem_freq_json   TEXT,
            flagged_json     TEXT,
            sentence_counts  TEXT
        );
        CREATE TABLE IF NOT EXISTS review_items (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            analysis_id INTEGER NOT NULL,
            item_type   TEXT NOT NULL,
            item_text   TEXT NOT NULL,
            suggestion  TEXT,
            grade       INTEGER,
            status      TEXT DEFAULT 'pending',
            notes       TEXT,
            created_at  TEXT,
            updated_at  TEXT
        );
        CREATE TABLE IF NOT EXISTS book_glossary (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            analysis_id INTEGER NOT NULL,
            word        TEXT NOT NULL,
            definition  TEXT,
            grade       INTEGER,
            source      TEXT DEFAULT 'analysis',
            status      TEXT DEFAULT 'draft',
            created_at  TEXT,
            updated_at  TEXT
        );
        CREATE TABLE IF NOT EXISTS reading_attempts (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            student_name    TEXT,
            grade           INTEGER,
            passage_text    TEXT,
            transcript      TEXT,
            engine          TEXT,
            strictness      TEXT DEFAULT 'gentle',
            score_json      TEXT,
            audio_path      TEXT,
            created_at      TEXT
        );
        CREATE INDEX IF NOT EXISTS idx_grade_words_grade ON grade_words(grade);
        CREATE INDEX IF NOT EXISTS idx_grade_words_word  ON grade_words(word);
        CREATE INDEX IF NOT EXISTS idx_grade_files_grade ON grade_files(grade);
        CREATE INDEX IF NOT EXISTS idx_wgm_grade          ON word_grade_map(first_grade);
        CREATE INDEX IF NOT EXISTS idx_review_analysis    ON review_items(analysis_id);
        CREATE INDEX IF NOT EXISTS idx_glossary_analysis  ON book_glossary(analysis_id);
        CREATE INDEX IF NOT EXISTS idx_reading_grade      ON reading_attempts(grade);
    ''')
    conn.commit()
    conn.close()

# ── OCR helpers ───────────────────────────────────────────────────────────────

_TAMIL_OCR_CONFIG      = '--oem 3 --psm 6 -c preserve_interword_spaces=1'
_TAMIL_OCR_CONFIG_PSM3 = '--oem 3 --psm 3 -c preserve_interword_spaces=1'
_TAMIL_OCR_CONFIG_PSM4 = '--oem 3 --psm 4 -c preserve_interword_spaces=1'

def _normalize_tamil_text(text):
    """Normalize Unicode Tamil text without changing valid letters."""
    if not text:
        return ''
    import unicodedata
    text = unicodedata.normalize('NFC', text)
    text = re.sub(r'[\u200B-\u200D\uFEFF]', '', text)
    text = re.sub(r'[\t\r\f\v]+', ' ', text)
    text = re.sub(r' *\n *', '\n', text)
    text = re.sub(r' {2,}', ' ', text)
    return text.strip()

def _preprocess_ocr_image(img):
    """
    Clean rendered PDF pages before Tamil OCR.

    tntextbooks.in 2024/2025 PDFs are large vector-renders (40-76 MB).
    Tamil script has fine curved strokes — the old threshold (>170) was
    erasing them. Use softer threshold (>128) to preserve fine detail.
    Skip upscaling for vector renders — they're already high resolution.
    Skip median filter — it blurs fine Tamil curves in vector output.
    """
    try:
        from PIL import ImageOps
        gray = ImageOps.grayscale(img)
        w, h = gray.size
        # Only upscale genuinely low-res pages (actual photo scans)
        if max(w, h) < 1500:
            gray = gray.resize((w * 2, h * 2))
        # Gentle contrast — preserve fine strokes
        gray = ImageOps.autocontrast(gray, cutoff=1)
        # Softer threshold: 128 instead of 170
        gray = gray.point(lambda px: 255 if px > 128 else 0)
        return gray
    except Exception:
        return img


def _tesseract_languages():
    """Prefer Tamil-only OCR; fall back gracefully if tam.traineddata is missing."""
    try:
        import pytesseract
        langs = set(pytesseract.get_languages(config=''))
        if 'tam' in langs:
            return 'tam'
        if 'eng' in langs:
            logging.getLogger('app').warning(
                'Tamil Tesseract language pack not found. Install tesseract-ocr-tam or tam.traineddata.'
            )
            return 'eng'
    except Exception:
        pass
    return 'tam'

def _fix_common_tamil_ocr_errors(text):
    """Conservative Tamil OCR cleanup; add verified textbook-specific fixes here."""
    text = _normalize_tamil_text(text)
    corrections = {
        'புதுரமயான': 'புதுமையான',
        'வடிவரமப்பு': 'வடிவமைப்பு',
        'பபாருள்': 'பொருள்',
        'மறறும்': 'மற்றும்',
        'குழநரதைகளின்': 'குழந்தைகளின்',
        'உைவியல்': 'உளவியல்',
        '்ாரநதை': 'சார்ந்த',
        'அணுகுமுரை': 'அணுகுமுறை',
        'புதுரமகள்': 'புதுமைகள்',
        'தைாஙகி': 'தாங்கி',
        'உஙகளுரடய': 'உங்களுடைய',
        'கைஙகளில்': 'கரங்களில்',
        'தைவழும்பபாழுது': 'தவழும்பொழுது',
        'பபருமிதைம்': 'பெருமிதம்',
        'தைதும்ப': 'ததும்ப',
        'நுரழவீரகள்': 'நுழைவீர்கள்',
        'நம்புகிபைாம்': 'நம்புகிறோம்',
    }
    for wrong, right in corrections.items():
        text = text.replace(wrong, right)
    return text

# ── Text extraction ───────────────────────────────────────────────────────────

def extract_text(filepath):
    """
    Multi-strategy Tamil text extraction.
    1. pdfminer page-by-page  — standard Unicode Tamil PDFs (~1s/page)
    2. pdfplumber page-by-page — custom/Type3 font encodings
    3. Tesseract OCR           — scanned PDFs OR CID-font PDFs (~3-5s/page)
       CID fonts: Tamil typeset with Shree-Lipi, e-Kalappai, ELCOT, etc.
       produce (cid:N) tokens — very common in Tamil Nadu government textbooks.
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.txt':
        for enc in ['utf-8', 'utf-16', 'latin-1']:
            try:
                with open(filepath, 'r', encoding=enc) as f:
                    return _normalize_tamil_text(f.read())
            except UnicodeDecodeError:
                continue
        return ''

    if ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            parts = []
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                    if cells:
                        parts.append(' '.join(cells))
            return _normalize_tamil_text('\n'.join(parts))
        except Exception as e:
            logging.getLogger('app').warning(f'DOCX extraction failed {os.path.basename(filepath)}: {e}')
            return ''

    if ext != '.pdf':
        return ''

    fname = os.path.basename(filepath)
    has_any_text   = False
    is_cid_encoded = False   # True = Tamil DTP font, needs OCR
    _tamil_found   = False   # True = actual Tamil Unicode chars found

    # Strategy 1: pdfminer page-by-page
    try:
        from pdfminer.layout import LAParams
        from pdfminer.pdfpage import PDFPage
        from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
        from pdfminer.converter import TextConverter
        import io as _io

        rsrcmgr = PDFResourceManager()
        parts = []
        with open(filepath, 'rb') as fh:
            for page_num, page in enumerate(PDFPage.get_pages(fh)):
                try:
                    buf = _io.StringIO()
                    device = TextConverter(rsrcmgr, buf, laparams=LAParams())
                    PDFPageInterpreter(rsrcmgr, device).process_page(page)
                    device.close()
                    parts.append(buf.getvalue())
                except Exception as pe:
                    logging.getLogger('app').debug(f'{fname} p{page_num}: {pe}')

        text = '\n'.join(parts)
        if re.search(r'[\u0B80-\u0BFF]{2,}', text):
            _tamil_found = True
            _stage('done', 'text extracted (pdfminer)')
            return _normalize_tamil_text(text)   # Unicode Tamil found — done

        # Detect CID-encoded fonts: (cid:N) tokens dominate output.
        # This is the signature of Shree-Lipi / e-Kalappai / ELCOT fonts
        # used heavily in Tamil Nadu government textbook PDFs.
        cid_count = len(re.findall(r'\(cid:\d+\)', text))
        non_cid   = len(re.sub(r'\(cid:\d+\)', '', text).split())
        if cid_count > 10 and cid_count > non_cid:
            is_cid_encoded = True
            logging.getLogger('app').info(
                f'{fname}: CID font encoding detected ({cid_count} CID tokens) '
                f'— routing to OCR'
            )
        else:
            # Detect TAM/Bamini/TSCII font encoding:
            # Tamil chars stored as Latin-1 Extended codepoints (0xA0-0xFF).
            # Common in older TN textbooks typeset with TAM, TAB, Bamini fonts.
            # Signature: >10% of printable chars fall in 0xA0-0xFF range.
            printable = sum(1 for c in text if c.isprintable() and c != ' ')
            latin_ext = sum(1 for c in text if 0xA0 <= ord(c) <= 0xFF)
            if printable > 20 and latin_ext / max(printable, 1) > 0.10:
                is_cid_encoded = True   # both CID and TAM need OCR
                logging.getLogger('app').info(
                    f'{fname}: TAM/Bamini font encoding detected '
                    f'({latin_ext} Latin-ext chars, {latin_ext/max(printable,1):.0%} of text) '
                    f'— routing to OCR'
                )
            else:
                has_any_text = bool(text.strip())

    except Exception as e:
        logging.getLogger('app').debug(f'pdfminer failed {fname}: {e}')

    # Strategy 2: pdfplumber — skip if already identified as CID-encoded
    if not is_cid_encoded:
        try:
            import pdfplumber
            parts2 = []
            with pdfplumber.open(filepath) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        t = page.extract_text()
                        if t:
                            parts2.append(t)
                    except Exception as pe:
                        logging.getLogger('app').debug(f'{fname} plumber p{page_num}: {pe}')

            text2 = '\n'.join(parts2)
            if re.search(r'[\u0B80-\u0BFF]{2,}', text2):
                _tamil_found = True
                _stage('done', 'text extracted (pdfplumber)')
                return _normalize_tamil_text(text2)

            cid2     = len(re.findall(r'\(cid:\d+\)', text2))
            non_cid2 = len(re.sub(r'\(cid:\d+\)', '', text2).split())
            if cid2 > 10 and cid2 > non_cid2:
                is_cid_encoded = True
            else:
                printable2 = sum(1 for c in text2 if c.isprintable() and c != ' ')
                latin_ext2 = sum(1 for c in text2 if 0xA0 <= ord(c) <= 0xFF)
                if printable2 > 20 and latin_ext2 / max(printable2, 1) > 0.10:
                    is_cid_encoded = True
                else:
                    has_any_text = has_any_text or bool(text2.strip())

        except Exception as e:
            logging.getLogger('app').debug(f'pdfplumber failed {fname}: {e}')

    # Strategy 3: Tesseract OCR
    # Triggered whenever pdfminer/pdfplumber did not return usable Unicode Tamil.
    # Some textbook PDFs contain only tiny English selectable text such as
    # watermarks/page labels plus full-page Tamil images. In that case
    # has_any_text is True, but Tamil text is still absent, so OCR must run.
    if True:
        try:
            import gc
            import tempfile
            import pytesseract
            from PIL import Image
            from pdf2image import convert_from_path, pdfinfo_from_path

            reason = 'CID font encoding' if is_cid_encoded else 'no usable Tamil Unicode text; trying OCR fallback'
            dpi = int(os.environ.get('TAMIL_ANALYZER_OCR_DPI', '300'))
            max_pages = int(os.environ.get('TAMIL_ANALYZER_OCR_MAX_PAGES', '0'))
            ocr_timeout = int(os.environ.get('TAMIL_ANALYZER_OCR_TIMEOUT', '90'))
            ocr_lang = _tesseract_languages()

            info = pdfinfo_from_path(filepath)
            total_pages = int(info.get('Pages', 0) or 0)
            if max_pages > 0:
                total_pages = min(total_pages, max_pages)

            _stage('ocr', f'starting OCR — {total_pages} pages ({reason})')
            logging.getLogger('app').info(
                f'{fname}: {reason} — running low-memory OCR page-by-page '
                f'({total_pages} pages, {dpi} DPI)…'
            )

            ocr_parts = []
            # Track if we're getting Tamil — if first 3 pages yield nothing,
            # retry with psm 3 (auto-layout) which handles mixed Tamil/image pages
            psm6_tamil_count = 0
            use_psm3 = False

            with tempfile.TemporaryDirectory(prefix='tamil_ocr_') as tmpdir:
                for page_no in range(1, total_pages + 1):
                    try:
                        _stage('ocr', f'page {page_no} of {total_pages} — OCR in progress')
                        paths = convert_from_path(
                            filepath,
                            dpi=dpi,
                            first_page=page_no,
                            last_page=page_no,
                            fmt='png',
                            grayscale=True,
                            thread_count=1,
                            output_folder=tmpdir,
                            paths_only=True,
                        )
                        if not paths:
                            continue

                        img_path = paths[0]
                        with Image.open(img_path) as img:
                            clean_img = _preprocess_ocr_image(img)

                            # Choose config: after 3 pages with no Tamil, switch to psm 3
                            ocr_cfg = _TAMIL_OCR_CONFIG_PSM3 if use_psm3 else _TAMIL_OCR_CONFIG
                            txt = pytesseract.image_to_string(
                                clean_img,
                                lang=ocr_lang,
                                config=ocr_cfg,
                                timeout=ocr_timeout,
                            )
                            ocr_parts.append(txt or '')

                            # After first 3 content pages, check if we're getting Tamil
                            if page_no == 3 and not use_psm3:
                                tamil_so_far = len(re.findall(
                                    r'[\u0B80-\u0BFF]{2,}', '\n'.join(ocr_parts)
                                ))
                                if tamil_so_far == 0:
                                    use_psm3 = True
                                    logging.getLogger('app').info(
                                        f'{fname}: psm 6 found no Tamil in first 3 pages, '
                                        f'switching to psm 3 (auto-layout)'
                                    )
                                    _stage('ocr', f'switching OCR mode — retrying with auto-layout')

                        try:
                            os.remove(img_path)
                        except OSError:
                            pass

                        if page_no == 1 or page_no % 5 == 0 or page_no == total_pages:
                            logging.getLogger('app').info(
                                f'{fname}: OCR progress {page_no}/{total_pages}'
                            )

                    except RuntimeError:
                        logging.getLogger('app').warning(
                            f'{fname}: OCR timed out on page {page_no}; continuing'
                        )
                    except Exception as oe:
                        logging.getLogger('app').warning(
                            f'{fname}: OCR failed on page {page_no}: {oe}'
                        )
                    finally:
                        gc.collect()

            ocr_text = _fix_common_tamil_ocr_errors('\n'.join(ocr_parts))
            if re.search(r'[\u0B80-\u0BFF]{2,}', ocr_text):
                n = len(re.findall(r'[\u0B80-\u0BFF]{2,}', ocr_text))
                logging.getLogger('app').info(
                    f'{fname}: OCR complete — {n} Tamil words found'
                )
                _stage('done', f'OCR complete — {n} Tamil words found')
                return _normalize_tamil_text(ocr_text)
            else:
                logging.getLogger('app').warning(
                    f'{fname}: OCR ran but found no Tamil text. '
                    f'Ensure Tamil pack is installed: sudo apt install tesseract-ocr-tam'
                )

        except ImportError as ie:
            logging.getLogger('app').warning(
                f'{fname}: OCR libraries not available ({ie}). '
                f'Run: pip install pytesseract pdf2image'
            )
        except Exception as e:
            logging.getLogger('app').warning(f'{fname}: OCR failed — {e}')

    if has_any_text and not is_cid_encoded:
        logging.getLogger('app').warning(
            f'{fname}: selectable text was present but no usable Tamil Unicode was found, '
            f'and OCR did not produce Tamil. This may be a non-Unicode Tamil font '
            f'(TSCII/Bamini/TAB) or a low-quality scan.'
        )
    return ''


# Sandhi liaison consonants that appear only at word boundaries, never as
# meaningful word endings: த், க், ச், ட், ப்
_SANDHI_LIAISON = re.compile(r'[தகசடப]்$')

# Common connective words that merge with the preceding word in some texts:
# வந்தபோது ↔ வந்த போது
_COMPOUND_JOINERS = [
    'போது','பொழுது','வரை','தோறும்','விட','என்று',
    'என்பது','இருந்தும்','கொண்டு','கொண்ட',
]

def _normalize_word(word):
    """
    Strip trailing sandhi liaison consonants and return the clean stem input.
    அவனைத் → அவனை,  அவனைக் → அவனை,  நன்மைக் → நன்மை
    Words ending in ன்/ம்/ர்/ல் etc. are NOT touched — those are real endings.
    """
    if _SANDHI_LIAISON.search(word) and len(word) > 2:
        return word[:-2]
    return word

def _split_compound(word):
    """
    Split common compound words written as one token into their parts.
    வந்தபோது → [வந்த, போது],  பார்க்கும்போது → [பார்க்கும், போது]
    Returns a list with one item if no split is found.
    """
    for j in _COMPOUND_JOINERS:
        if word.endswith(j) and len(word) > len(j) + 1:
            prefix = word[:-len(j)]
            if len(prefix) >= 2:
                return [prefix, j]
    return [word]

def tokenize_tamil(text):
    """
    Extract Tamil words from text with three normalization passes:
      1. Extract Tamil Unicode tokens (U+0B80–U+0BFF), length >= 2
      2. Split compound words (வந்தபோது → வந்த + போது)
      3. Strip trailing sandhi liaison consonants (அவனைத் → அவனை)

    The final filter is important for noisy textbook PDFs: compound splitting or
    OCR cleanup can sometimes leave a single combining mark / empty token.
    """
    raw = re.findall(r'[\u0B80-\u0BFF]{2,}', text or '')
    result = []
    for word in raw:
        for part in _split_compound(word):
            part = _normalize_word(part).strip()
            if len(part) >= 2 and re.search(r'[\u0B80-\u0BFF]', part):
                result.append(part)
    return result

# ── Routes ────────────────────────────────────────────────────────────────────

@app.route('/')
def index(): return render_template('index.html')

@app.route('/api/grades')
def get_grades():
    conn = get_db()
    metas = conn.execute('SELECT * FROM grade_meta ORDER BY grade').fetchall()
    files = conn.execute('SELECT * FROM grade_files ORDER BY grade, processed_at').fetchall()
    conn.close()
    files_by_grade = {}
    for f in files:
        g = f['grade']
        if g not in files_by_grade: files_by_grade[g] = []
        files_by_grade[g].append(dict(f))
    result = []
    for m in metas:
        d = dict(m)
        d['files'] = files_by_grade.get(m['grade'], [])
        result.append(d)
    return jsonify(result)

@app.route('/api/stemmer_info')
def stemmer_info():
    return jsonify({'snowball_available': SNOWBALL_AVAILABLE})

@app.route('/api/stem_demo', methods=['POST'])
def stem_demo():
    words = request.json.get('words', [])
    result = []
    for w in words[:30]:
        w = w.strip()
        if re.search(r'[\u0B80-\u0BFF]', w):
            result.append({'original': w, 'stem': get_stem(w)})
    return jsonify(result)

# ── Processing lock prevents concurrent SQLite writes ────────────────────────
_db_write_lock = threading.Lock()

def _compute_file_hash(filepath):
    h = hashlib.md5()
    with open(filepath, 'rb') as f:
        for chunk in iter(lambda: f.read(65536), b''): h.update(chunk)
    return h.hexdigest()

def _extract_stems_parallel(filepath):
    """CPU-bound: extract text and compute stems. Safe to run in threads."""
    text = extract_text(filepath)
    raw_words = tokenize_tamil(text)
    if not raw_words:
        return None, None, None, None
    stems_set = set()
    for w in raw_words:
        s = get_stem(w)
        if s and len(s) >= 2 and re.search(r'[\u0B80-\u0BFF]', s):
            stems_set.add(s)
    stems = list(stems_set)
    sc = sentence_word_counts(text)
    ss = sentence_stats(sc)
    return stems, len(set(raw_words)), ss, text

def _rebuild_grade_aggregate(conn, grade):
    """Recompute grade_meta from all grade_files rows for this grade."""
    rows = conn.execute(
        'SELECT * FROM grade_files WHERE grade = ?', (grade,)
    ).fetchall()
    if not rows:
        conn.execute('DELETE FROM grade_meta WHERE grade = ?', (grade,))
        return
    file_count  = len(rows)
    total_raw   = sum(r['raw_count'] for r in rows)
    # Sentence stats: take the max across all files (most demanding)
    sent_max    = max(r['sent_max'] for r in rows)
    # Weighted average for sent_avg
    total_sents = sum(r['sent_total'] for r in rows)
    sent_avg    = round(
        sum(r['sent_avg'] * r['sent_total'] for r in rows) / max(total_sents, 1), 1
    )
    # Approximate median as average of file medians
    sent_median = int(sum(r['sent_median'] for r in rows) / len(rows))
    # word_count = total unique stems across all files for this grade
    wc = conn.execute(
        'SELECT COUNT(*) FROM grade_words WHERE grade = ?', (grade,)
    ).fetchone()[0]
    conn.execute('''
        INSERT OR REPLACE INTO grade_meta
          (grade, file_count, word_count, raw_count,
           sent_max, sent_avg, sent_median, sent_total, updated_at)
        VALUES (?,?,?,?,?,?,?,?,?)
    ''', (grade, file_count, wc, total_raw,
          sent_max, sent_avg, sent_median, total_sents,
          datetime.datetime.now().isoformat()))

# Pre-loaded hash cache — populated by _load_all_hashes() before batch scans
# so individual file checks don't each need a DB roundtrip
_hash_cache = {}   # {norm_path: file_hash}
_hash_cache_lock = threading.Lock()

def _load_all_hashes():
    """Load all known file hashes into memory in one query. Call before a batch scan."""
    global _hash_cache
    conn = get_db()
    rows = conn.execute('SELECT filepath, file_hash FROM grade_files').fetchall()
    conn.close()
    with _hash_cache_lock:
        _hash_cache = {r['filepath']: r['file_hash'] for r in rows}

def _process_grade_file(filepath, grade, source='folder'):
    """
    Process one file for a grade. Supports multiple files per grade.
    - Skips if file hash unchanged (already in DB or in _hash_cache).
    - Adds only NEW stems (doesn't delete existing stems from other files).
    - Rebuilds grade aggregate stats after each file.
    - Thread-safe: text extraction runs freely, DB writes are locked.
    """
    fhash     = _compute_file_hash(filepath)
    filename  = os.path.basename(filepath)
    norm_path = os.path.normpath(os.path.abspath(filepath))

    # Check in-memory cache first (populated before batch scan)
    with _hash_cache_lock:
        cached = _hash_cache.get(norm_path)
    if cached == fhash:
        return {'skipped': True, 'filename': filename, 'grade': grade,
                'reason': 'unchanged'}

    # Fallback: check DB directly (for individual uploads)
    if cached is None:
        with _db_write_lock:
            conn = get_db()
            existing = conn.execute(
                'SELECT file_hash FROM grade_files WHERE filepath = ?', (norm_path,)
            ).fetchone()
            conn.close()
        if existing and existing['file_hash'] == fhash:
            with _hash_cache_lock:
                _hash_cache[norm_path] = fhash
            return {'skipped': True, 'filename': filename, 'grade': grade,
                    'reason': 'unchanged'}

    # CPU-heavy extraction — runs outside the lock
    try:
        import folder_watcher as _fwmod
        _fwmod.update_file_stage(filename, 'extracting', 'starting text extraction')
    except Exception:
        pass
    stems, raw_unique, ss, _ = _extract_stems_parallel(filepath)
    if stems is None:
        return {'error': 'No Tamil text found. Either: (1) scanned image PDF — needs OCR, or (2) non-Unicode Tamil font (TSCII/Bamini) — convert to Unicode PDF.',
                'filename': filename, 'grade': grade}

    now = datetime.datetime.now().isoformat()

    with _db_write_lock:
        conn = get_db()

        # Upsert file record
        conn.execute('''
            INSERT INTO grade_files
              (grade, filepath, filename, file_hash, word_count, raw_count,
               sent_max, sent_avg, sent_median, sent_total, source, processed_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?)
            ON CONFLICT(filepath) DO UPDATE SET
              file_hash=excluded.file_hash, word_count=excluded.word_count,
              raw_count=excluded.raw_count, sent_max=excluded.sent_max,
              sent_avg=excluded.sent_avg, sent_median=excluded.sent_median,
              sent_total=excluded.sent_total, processed_at=excluded.processed_at
        ''', (grade, norm_path, filename, fhash, len(stems), raw_unique,
              ss['max'], ss['avg'], ss['median'], ss['total_sentences'],
              source, now))

        # Batch insert stems — INSERT OR IGNORE keeps existing words from
        # other files for the same grade. executemany with a list is fast;
        # for very large files (>50 k stems) we chunk to avoid locking too long.
        CHUNK = 5000
        for i in range(0, len(stems), CHUNK):
            conn.executemany(
                'INSERT OR IGNORE INTO grade_words (grade, word) VALUES (?, ?)',
                [(grade, s) for s in stems[i:i+CHUNK]]
            )

        # Update word→grade map in the same chunked loop
        for i in range(0, len(stems), CHUNK):
            conn.executemany('''
                INSERT INTO word_grade_map (stem, first_grade) VALUES (?, ?)
                ON CONFLICT(stem) DO UPDATE
                  SET first_grade = MIN(first_grade, excluded.first_grade)
            ''', [(s, grade) for s in stems[i:i+CHUNK]])

        _rebuild_grade_aggregate(conn, grade)
        conn.commit()
        conn.close()

    # Update in-memory cache so subsequent checks are instant
    with _hash_cache_lock:
        _hash_cache[norm_path] = fhash

    # Clear stage tracking
    try:
        import folder_watcher as _fwmod
        _fwmod.clear_file_stage(filename)
    except Exception:
        pass

    return {'grade': grade, 'word_count': len(stems), 'raw_count': raw_unique,
            'filename': filename, 'file_hash': fhash,
            'reduction_pct': round((1 - len(stems)/max(raw_unique,1))*100, 1),
            'sent_max': ss['max'], 'sent_avg': ss['avg']}

@app.route('/api/upload_grade', methods=['POST'])
def upload_grade():
    grade = int(request.form.get('grade', 0))
    if not 1 <= grade <= 12:
        return jsonify({'error': 'Invalid grade'}), 400
    files = request.files.getlist('file')
    if not files: return jsonify({'error': 'No files'}), 400

    results = []
    for file in files:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], f'grade_{grade}_{filename}')
        file.save(filepath)
        result = _process_grade_file(filepath, grade, source='upload')
        try: os.remove(filepath)
        except: pass
        results.append(result)

    errors = [r for r in results if 'error' in r]
    ok     = [r for r in results if 'error' not in r and not r.get('skipped')]
    # Return aggregated grade info
    conn = get_db()
    meta = conn.execute('SELECT * FROM grade_meta WHERE grade = ?', (grade,)).fetchone()
    file_rows = conn.execute('SELECT * FROM grade_files WHERE grade = ? ORDER BY processed_at DESC', (grade,)).fetchall()
    conn.close()
    return jsonify({
        'grade': grade,
        'files_processed': len(ok),
        'files_skipped': len([r for r in results if r.get('skipped')]),
        'errors': errors,
        'grade_word_count': dict(meta)['word_count'] if meta else 0,
        'grade_file_count': dict(meta)['file_count'] if meta else 0,
        'files': [dict(r) for r in file_rows],
    })

@app.route('/api/delete_grade', methods=['POST'])
def delete_grade():
    data = request.json or {}
    grade    = data.get('grade')
    filepath = data.get('filepath')   # optional: delete just one file
    conn = get_db()
    conn.execute('PRAGMA journal_mode=WAL')
    if filepath:
        # Remove one specific file from a grade
        norm = os.path.normpath(os.path.abspath(filepath))
        conn.execute('DELETE FROM grade_files WHERE filepath = ?', (norm,))
        # Rebuild grade_words from remaining files (slow but correct)
        conn.execute('DELETE FROM grade_words WHERE grade = ?', (grade,))
        remaining = conn.execute(
            'SELECT filepath FROM grade_files WHERE grade = ?', (grade,)
        ).fetchall()
        conn.commit()
        conn.close()
        # Reprocess remaining files to rebuild vocabulary
        for row in remaining:
            _process_grade_file(row['filepath'], grade, source='folder')
        # Also fix word_grade_map
        _rebuild_word_grade_map()
    else:
        # Remove entire grade
        conn.execute('DELETE FROM grade_words WHERE grade = ?', (grade,))
        conn.execute('DELETE FROM grade_meta WHERE grade = ?', (grade,))
        conn.execute('DELETE FROM grade_files WHERE grade = ?', (grade,))
        # Fix word_grade_map for stems that may have belonged to this grade
        conn.execute('DELETE FROM word_grade_map')
        conn.commit()
        conn.close()
        _rebuild_word_grade_map()
    return jsonify({'ok': True})

def _rebuild_word_grade_map():
    """Full rebuild of word_grade_map from grade_words. Called after deletions."""
    with _db_write_lock:
        conn = get_db()
        conn.execute('DELETE FROM word_grade_map')
        conn.execute('''
            INSERT INTO word_grade_map (stem, first_grade)
            SELECT word, MIN(grade) FROM grade_words GROUP BY word
        ''')
        conn.commit()
        conn.close()

@app.route('/api/extract', methods=['POST'])
def extract_for_review():
    file = request.files.get('file')
    if not file: return jsonify({'error': 'No file'}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], f'target_{filename}')
    file.save(filepath)
    text = extract_text(filepath)
    all_words = tokenize_tamil(text)

    # Capture sentence counts from target book
    target_sent_counts = sentence_word_counts(text)
    os.remove(filepath)

    if not all_words: return jsonify({'error': 'No Tamil words found.'}), 400

    total_words = len(all_words)
    unique_raw = len(set(all_words))
    stem_to_original = {}; stem_freq = {}; all_stems = []
    for w in all_words:
        s = get_stem(w); all_stems.append(s)
        if s not in stem_to_original: stem_to_original[s] = w
        stem_freq[s] = stem_freq.get(s, 0) + 1

    unique_stems = set(all_stems)
    conn = get_db()
    rows = conn.execute('SELECT word FROM grade_words').fetchall()
    conn.close()
    grade_vocab_union = {r['word'] for r in rows}
    if not grade_vocab_union:
        return jsonify({'error': 'No school books uploaded yet.'}), 400

    flagged = detect_proper_nouns(stem_freq, grade_vocab_union)
    flagged_list = []
    for stem, reasons in sorted(flagged.items(), key=lambda x: -stem_freq.get(x[0],0)):
        flagged_list.append({
            'stem': stem, 'display': stem_to_original.get(stem, stem),
            'freq': stem_freq.get(stem, 0), 'reasons': reasons,
            'suggested': 'proper_noun' if any(
                r in reasons for r in ['deity/god name','place-name suffix',
                                       'person-name suffix','foreign name pattern','loan-word starter']
            ) else 'rare_unknown'
        })

    conn = get_db()
    conn.execute('DELETE FROM pending_extractions')
    conn.execute('''
        INSERT INTO pending_extractions
          (book_name, created_at, total_words, unique_words, unique_stems,
           stem_to_original, stem_freq_json, flagged_json, sentence_counts)
        VALUES (?,?,?,?,?,?,?,?,?)
    ''', (filename, datetime.datetime.now().isoformat(),
          total_words, unique_raw, len(unique_stems),
          json.dumps(stem_to_original), json.dumps(stem_freq),
          json.dumps(flagged), json.dumps(target_sent_counts)))
    # Store raw text for analytics (trimmed to 500 KB)
    try:
        conn.execute("ALTER TABLE pending_extractions ADD COLUMN raw_text TEXT")
    except: pass
    conn.execute("UPDATE pending_extractions SET raw_text = ? WHERE id = (SELECT last_insert_rowid())",
                 (text[:500000],))
    conn.commit()
    pending_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
    conn.close()

    return jsonify({
        'pending_id': pending_id, 'book_name': filename,
        'total_words': total_words, 'unique_words': unique_raw,
        'unique_stems': len(unique_stems),
        'flagged_count': len(flagged_list), 'flagged': flagged_list
    })


def _run_single_analysis(pending_id, confirmed_proper_stems=None):
    """
    Run the full analysis for a pending extraction without going through HTTP.
    Used by batch_analyze to skip proper-noun review.
    Returns the same dict that /api/analyze would return as JSON.
    """
    confirmed_proper = set(confirmed_proper_stems or [])

    conn = get_db()
    row  = conn.execute('SELECT * FROM pending_extractions WHERE id = ?',
                        (pending_id,)).fetchone()
    conn.close()
    if not row:
        return {'error': f'Pending extraction {pending_id} not found.'}

    # Build a mock request body and call analyze() by re-creating its logic
    # We do this by directly importing the request context helper
    from flask import Flask
    with app.test_request_context(
            '/api/analyze', method='POST',
            json={'pending_id': pending_id, 'proper_noun_stems': list(confirmed_proper)},
            content_type='application/json'):
        resp = analyze()
        # resp may be a Response or a (data, code, headers) tuple
        if isinstance(resp, tuple):
            data, *_ = resp
            if hasattr(data, 'get_json'):
                return data.get_json() or {}
            return {}
        if hasattr(resp, 'get_json'):
            return resp.get_json() or {}
        return {}

@app.route('/api/analyze', methods=['POST'])
def analyze():
    data = request.json
    pending_id = data.get('pending_id')
    confirmed_proper = set(data.get('proper_noun_stems', []))

    conn = get_db()
    row = conn.execute('SELECT * FROM pending_extractions WHERE id = ?', (pending_id,)).fetchone()
    conn.close()
    if not row: return jsonify({'error': 'Pending extraction not found.'}), 400

    book_name        = row['book_name']
    total_words      = row['total_words']
    unique_words     = row['unique_words']
    unique_stems_cnt = row['unique_stems']
    stem_to_original = json.loads(row['stem_to_original'])
    stem_freq        = json.loads(row['stem_freq_json'])
    target_sent_counts = json.loads(row['sentence_counts'] or '[]')

    all_stems = set(stem_freq.keys())

    # Load grade vocab and grade sentence stats
    conn = get_db()
    grade_rows  = conn.execute('SELECT grade, word FROM grade_words ORDER BY grade').fetchall()
    grade_metas = conn.execute('SELECT * FROM grade_meta ORDER BY grade').fetchall()
    conn.close()

    grade_vocab = {}
    for r in grade_rows:
        g = r['grade']
        if g not in grade_vocab: grade_vocab[g] = set()
        grade_vocab[g].add(r['word'])

    grade_sent_stats = {r['grade']: dict(r) for r in grade_metas}
    available_grades = sorted(grade_vocab.keys())
    if not available_grades: return jsonify({'error': 'No grade books in database.'}), 400

    effective_total = len(all_stems)

    # ── Per-grade new-word analysis ───────────────────────────────────────────
    # "New words for Std N" = words in target book that are UNKNOWN up to Std N-1
    # but KNOWN at Std N (i.e. first introduced at that standard).
    # This tells: if a student just finished Std N, these are the words
    # they would have newly learned that help them read this book.

    results = []
    cumulative_prev = set()
    cumulative = set()

    for g in range(1, 13):
        if g in grade_vocab: cumulative |= grade_vocab[g]
        if g not in available_grades: continue

        vocab_known = all_stems & cumulative
        known = vocab_known | confirmed_proper
        unknown = all_stems - known

        # New words at this grade = words known now but unknown up to previous grade
        prev_known = (all_stems & cumulative_prev) | confirmed_proper
        new_at_this_grade = known - prev_known
        new_word_list = sorted([stem_to_original.get(s,s) for s in new_at_this_grade])

        comprehension = round(len(known)/effective_total*100, 1) if effective_total else 0.0
        new_pct       = round(len(new_at_this_grade)/effective_total*100, 1) if effective_total else 0.0

        # Sentence complexity for this grade
        gsm = grade_sent_stats.get(g, {})
        grade_max = gsm.get('sent_max', 0)
        grade_avg = gsm.get('sent_avg', 0.0)

        # How many target-book sentences exceed this grade's max sentence length?
        sentences_over = sum(1 for c in target_sent_counts if c > grade_max) if grade_max > 0 else 0
        pct_over = round(sentences_over / len(target_sent_counts) * 100, 1) if target_sent_counts else 0.0

        new_to_student_pct = round(len(unknown)/effective_total*100, 1) if effective_total else 0.0
        results.append({
            'grade': g,
            'total_unique_book_words': effective_total,
            'known_words': len(known),        # words in book already known by student (Std 1–N)
            'new_words':   len(unknown),      # words in book new to student (not yet learned)
            'known_pct':   comprehension,     # % known
            'new_pct':     new_to_student_pct,# % new to student
            'new_word_list': sorted([stem_to_original.get(s,s) for s in unknown])[:500],
            # Words first introduced at this grade (for Section 3 word lists)
            'new_at_grade': len(new_at_this_grade),
            'new_at_grade_pct': new_pct,
            'new_at_grade_list': new_word_list,
            # Aliases for backward compat
            'found_words':    len(known),
            'not_found_words': len(unknown),
            'found_pct':      comprehension,
            'not_found_pct':  new_to_student_pct,
            'comprehension_pct': comprehension,
            'unknown_word_list': sorted([stem_to_original.get(s,s) for s in unknown])[:500],
            # Sentence stats
            'grade_sent_max': grade_max,
            'grade_sent_avg': grade_avg,
            'target_sentences_over_max': sentences_over,
            'target_pct_over_max': pct_over,
        })

        cumulative_prev = set(cumulative)  # snapshot before next grade adds words

    first_readable = next((r for r in results if r['comprehension_pct'] >= 80), None)
    best_grade = first_readable['grade'] if first_readable else None

    distribution = []
    assigned_stems = set()
    for r in results:
        stems_here = []
        surface_words = set(r.get('new_at_grade_list') or [])
        for stem, surface in stem_to_original.items():
            if surface in surface_words:
                stems_here.append(stem)
        # Fall back to count/list from the result if surface duplicates make
        # reverse mapping imperfect. Display uses the original result list.
        words_here = r.get('new_at_grade_list') or []
        assigned_stems.update(stems_here)
        distribution.append({
            'grade': r['grade'],
            'word_count': r.get('new_at_grade', len(words_here)),
            'word_pct': r.get('new_at_grade_pct', round(len(words_here) / effective_total * 100, 1) if effective_total else 0),
            'words': words_here[:500],
        })
    unknown_after_all = set(results[-1].get('unknown_word_list', [])) if results else set()
    distribution.append({
        'grade': None,
        'label': 'Not found in any class',
        'word_count': len(unknown_after_all),
        'word_pct': round(len(unknown_after_all) / effective_total * 100, 1) if effective_total else 0,
        'words': sorted(unknown_after_all)[:500],
    })

    proper_noun_list = sorted([stem_to_original.get(s,s) for s in confirmed_proper])

    # Target book sentence stats
    tss = sentence_stats(target_sent_counts)

    # Run extended analytics (all free, local computation)
    try:
        conn2 = get_db()
        raw_text_row = conn2.execute(
            "SELECT raw_text FROM pending_extractions WHERE id = ?", (pending_id,)
        ).fetchone()
        conn2.close()
        raw_text = raw_text_row['raw_text'] if raw_text_row and raw_text_row['raw_text'] else ''
    except: raw_text = ''

    best_comp = next((r['comprehension_pct'] for r in results
                      if r['comprehension_pct'] >= 80), results[-1]['comprehension_pct'] if results else 0)
    best_r    = next((r for r in results if r['comprehension_pct'] >= 80), results[-1] if results else {})
    analytics_data = _analytics.full_analytics(
        raw_text,
        stem_fn           = get_stem,
        comprehension_pct = best_comp,
        sent_avg          = tss.get('avg', 0),
        sent_max_grade    = best_r.get('grade_sent_max', 0),
    )

    # Optional meaning-level analysis. Separate add-on layer; it does not change
    # existing readability scoring. It uses data/meaning_kb when built.
    try:
        target_grade_for_meaning = best_grade or (best_r.get('grade') if best_r else 12) or 12
        meaning_data = _meaning_kb.analyze_text_meaning(
            raw_text, int(target_grade_for_meaning), 'data/meaning_kb',
            tokenize_fn=tokenize_tamil, stem_fn=get_stem, limit=300
        )
    except Exception as e:
        meaning_data = {'enabled': False, 'error': str(e)}

    # v11 suitability/adaptation report. This is an add-on JSON blob; the
    # legacy results_json remains unchanged for backward compatibility.
    try:
        suitability_data = _suitability.build_suitability_report(
            raw_text=raw_text,
            results=results,
            target_sentence_counts=target_sent_counts,
            meaning=meaning_data,
            tokenize_fn=tokenize_tamil,
            stem_fn=get_stem,
            kb_dir='data/meaning_kb',
        )
        _suitability.save_analysis_cache(suitability_data, book_name, raw_text, 'data/cache')
        _suitability.update_books_index(suitability_data, book_name, 'data/books_index.json')
    except Exception as e:
        suitability_data = {'enabled': False, 'error': str(e)}

    # v27 offline intelligence layer: no live AI/API usage.
    try:
        v27_data = _v27.build_offline_intelligence(
            raw_text=raw_text,
            results=results,
            target_sentence_counts=target_sent_counts,
            meaning=meaning_data,
            suitability=suitability_data,
            kb_dir='data/meaning_kb',
        )
    except Exception as e:
        v27_data = {'enabled': False, 'error': str(e)}

    conn = get_db()
    try:
        conn.execute("ALTER TABLE analyses ADD COLUMN analytics_json TEXT")
    except: pass
    try:
        conn.execute("ALTER TABLE analyses ADD COLUMN meaning_json TEXT")
    except: pass
    try:
        conn.execute("ALTER TABLE analyses ADD COLUMN suitability_json TEXT")
    except: pass
    try:
        conn.execute("ALTER TABLE analyses ADD COLUMN v27_json TEXT")
    except: pass
    conn.execute('''
        INSERT INTO analyses
          (book_name, analyzed_at, total_words, unique_words, unique_stems,
           proper_nouns, sentence_json, results_json, analytics_json, meaning_json, suitability_json, v27_json)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?)
    ''', (book_name, datetime.datetime.now().isoformat(),
          total_words, unique_words, unique_stems_cnt,
          json.dumps(proper_noun_list),
          json.dumps({'target': tss, 'target_counts': target_sent_counts[:2000]}),
          json.dumps(results),
          json.dumps(analytics_data),
          json.dumps(meaning_data),
          json.dumps(suitability_data),
          json.dumps(v27_data)))
    conn.commit()
    analysis_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
    conn.close()

    return jsonify({
        'analysis_id': analysis_id,
        'book_name': book_name,
        'total_words': total_words,
        'unique_words': unique_words,
        'unique_stems': unique_stems_cnt,
        'proper_nouns_excluded': len(confirmed_proper),
        'proper_noun_list': proper_noun_list,
        'stems_analyzed': effective_total,
        'best_grade': best_grade,
        'target_sentence_stats': tss,
        'word_distribution': distribution,
        'analytics': analytics_data,
        'meaning': meaning_data,
        'suitability': suitability_data,
        'v27': v27_data,
        'results': results
    })


@app.route('/api/meaning/status')
def meaning_status():
    kb = _meaning_kb.load_kb('data/meaning_kb')
    if not kb:
        return jsonify({'built': False, 'message': 'Meaning knowledge base not built yet.'})
    return jsonify({'built': True, 'metadata': kb.get('metadata', {})})

@app.route('/api/meaning/build', methods=['POST'])
def meaning_build_api():
    data = request.json or {}
    full = bool(data.get('full_rebuild', True))
    try:
        meta = _meaning_kb.build_from_existing_db(
            DB_PATH, 'data/meaning_kb',
            extract_text_fn=extract_text,
            tokenize_fn=tokenize_tamil,
            stem_fn=get_stem,
            full_rebuild=full,
        )
        return jsonify({'ok': True, 'metadata': meta})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/api/books/compare')
def books_compare():
    names = request.args.getlist('book')
    try:
        return jsonify(_suitability.compare_books('data/books_index.json', names or None))
    except Exception as e:
        return jsonify({'error': str(e), 'books': []}), 500

@app.route('/api/books/index')
def books_index():
    try:
        p = os.path.join('data', 'books_index.json')
        if not os.path.exists(p):
            return jsonify([])
        with open(p, 'r', encoding='utf-8') as f:
            return jsonify(json.load(f))
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ── Textbook Auto Importer ───────────────────────────────────────────────────

@app.route('/api/importer/sources', methods=['GET'])
def importer_sources_get():
    return jsonify(_importer.load_sources())

@app.route('/api/importer/sources', methods=['POST'])
def importer_sources_post():
    data = request.json or {}
    if not data.get('url'):
        return jsonify({'error': 'Source URL is required'}), 400
    source = _importer.add_source({
        'name': data.get('name') or 'Textbook Source',
        'url': data.get('url'),
        'board': data.get('board') or '',
        'language': data.get('language') or 'Tamil',
        'tamil_only': data.get('tamil_only', True),
        'exclude_english': data.get('exclude_english', True),
        'class': data.get('class') or None,
        'active': data.get('active', True),
    })
    return jsonify({'ok': True, 'source': source})

@app.route('/api/importer/scan', methods=['POST'])
def importer_scan():
    data = request.json or {}
    source = data.get('source') or {}
    if not source.get('url'):
        sid = data.get('source_id')
        source = next((s for s in _importer.load_sources() if s.get('id') == sid), None)
    if not source or not source.get('url'):
        return jsonify({'error': 'Valid source or source_id is required'}), 400
    try:
        result = _importer.scan_source(source, depth=int(data.get('depth', 1)), max_pages=int(data.get('max_pages', 30)))
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/importer/recent')
def importer_recent():
    return jsonify(_importer.recent_discovered())


_IMPORT_JOBS = {}
_IMPORT_JOBS_LOCK = threading.Lock()

def _set_import_job(job_id, **updates):
    with _IMPORT_JOBS_LOCK:
        job = _IMPORT_JOBS.setdefault(job_id, {})
        job.update(updates)
        job["updated_at"] = time.time()
        return dict(job)

def _get_import_job(job_id):
    with _IMPORT_JOBS_LOCK:
        return dict(_IMPORT_JOBS.get(job_id, {}))

def _run_import_job(job_id, items, rebuild_meaning, workers):
    try:
        _set_import_job(job_id, status="running", phase="download", message="Starting downloads", total=len(items), downloaded=0, processed=0, results=[])
        _load_all_hashes()
        def progress(evt):
            updates = {"phase": evt.get("phase"), "message": evt.get("message")}
            if "downloaded" in evt: updates["downloaded"] = evt.get("downloaded")
            if "processed" in evt: updates["processed"] = evt.get("processed")
            if "total" in evt: updates["phase_total"] = evt.get("total")
            if evt.get("last"):
                job = _get_import_job(job_id)
                recent = job.get("recent", [])[-10:]
                recent.append(evt.get("last"))
                updates["recent"] = recent
            _set_import_job(job_id, **updates)
        result = _importer.download_items_parallel(items, process_fn=_process_grade_file, max_workers=workers, progress=progress)
        _set_import_job(job_id, results=result.get("results", []), downloaded=result.get("downloaded", 0), skipped=result.get("skipped", 0), failed=result.get("failed", 0))
        if rebuild_meaning:
            _set_import_job(job_id, phase="meaning", message="Rebuilding meaning data")
            try:
                meta = _meaning_kb.build_from_existing_db(DB_PATH, "data/meaning_kb", extract_text_fn=extract_text, tokenize_fn=tokenize_tamil, stem_fn=get_stem, full_rebuild=True)
                result["meaning_rebuilt"] = True
                result["meaning_metadata"] = meta
            except Exception as kb_err:
                result["meaning_rebuilt"] = False
                result["meaning_error"] = str(kb_err)
        else:
            result["meaning_rebuilt"] = False
        downloaded = result.get("downloaded", 0)
        skipped = result.get("skipped", 0)
        failed = result.get("failed", 0)
        processed = len([r for r in result.get("results", []) if r.get("processed")])
        _set_import_job(job_id, status="done", phase="done", downloaded=downloaded, skipped=skipped, failed=failed, processed=processed, message=f"Imported complete: {downloaded} downloaded, {skipped} skipped existing/excluded, {failed} failed, {processed} processed", result=result)
    except Exception as e:
        _set_import_job(job_id, status="error", phase="error", message=str(e), error=str(e))

@app.route("/api/importer/download_async", methods=["POST"])
def importer_download_async():
    data = request.json or {}
    items = data.get("items") or []
    if not items:
        return jsonify({"error": "No items selected"}), 400
    rebuild_meaning = bool(data.get("rebuild_meaning", True))
    workers = int(data.get("workers") or 3)
    workers = max(1, min(workers, 6))
    job_id = uuid.uuid4().hex[:12]
    _set_import_job(job_id, id=job_id, status="queued", phase="queued", message="Queued", total=len(items), downloaded=0, processed=0, created_at=time.time())
    t = threading.Thread(target=_run_import_job, args=(job_id, items, rebuild_meaning, workers), daemon=True)
    t.start()
    return jsonify({"ok": True, "job_id": job_id})

@app.route("/api/importer/download_status/<job_id>")
def importer_download_status(job_id):
    job = _get_import_job(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    return jsonify(job)

@app.route('/api/importer/download', methods=['POST'])
def importer_download():
    data = request.json or {}
    items = data.get('items') or []
    if not items:
        return jsonify({'error': 'No items selected'}), 400
    limit = data.get('limit')
    rebuild_meaning = bool(data.get('rebuild_meaning', True))
    try:
        _load_all_hashes()
        result = _importer.download_items(items, process_fn=_process_grade_file, limit=int(limit) if limit else None)

        # One-click import flow: after textbooks are imported into the existing
        # grade database, refresh the separate meaning KB. This does not change
        # the existing grade database schema or old analysis workflow.
        if rebuild_meaning:
            try:
                meta = _meaning_kb.build_from_existing_db(
                    DB_PATH, 'data/meaning_kb',
                    extract_text_fn=extract_text,
                    tokenize_fn=tokenize_tamil,
                    stem_fn=get_stem,
                    full_rebuild=True,
                )
                result['meaning_rebuilt'] = True
                result['meaning_metadata'] = meta
            except Exception as kb_err:
                result['meaning_rebuilt'] = False
                result['meaning_error'] = str(kb_err)
        else:
            result['meaning_rebuilt'] = False
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/history')
def get_history():
    conn = get_db()
    rows = conn.execute(
        'SELECT id, book_name, analyzed_at, total_words, unique_words, unique_stems '
        'FROM analyses ORDER BY id DESC LIMIT 50').fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/history/<int:analysis_id>')
def get_analysis(analysis_id):
    conn = get_db()
    row = conn.execute('SELECT * FROM analyses WHERE id = ?', (analysis_id,)).fetchone()
    conn.close()
    if not row: return jsonify({'error': 'Not found'}), 404
    d = dict(row)
    d['results'] = json.loads(d['results_json'])
    d['proper_noun_list'] = json.loads(d.get('proper_nouns') or '[]')
    d['sentence_data']    = json.loads(d.get('sentence_json') or '{}')
    d['meaning']          = json.loads(d.get('meaning_json') or 'null') if 'meaning_json' in d else None
    d['suitability']      = json.loads(d.get('suitability_json') or 'null') if 'suitability_json' in d else None
    del d['results_json']
    return jsonify(d)

def _load_analysis_payload(analysis_id):
    conn = get_db()
    row = conn.execute('SELECT * FROM analyses WHERE id = ?', (analysis_id,)).fetchone()
    conn.close()
    if not row:
        return None
    d = dict(row)
    d['results'] = json.loads(d.get('results_json') or '[]')
    d['proper_nouns'] = json.loads(d.get('proper_nouns') or '[]')
    d['sentence_data'] = json.loads(d.get('sentence_json') or '{}')
    d['meaning'] = json.loads(d.get('meaning_json') or 'null') if 'meaning_json' in d else None
    d['suitability'] = json.loads(d.get('suitability_json') or 'null') if 'suitability_json' in d else None
    return d

def _review_candidates_from_analysis(payload):
    items = []
    analysis_id = payload['id']
    suitability = payload.get('suitability') or {}
    rec_grade = suitability.get('recommended_grade') or next(
        (r.get('grade') for r in payload.get('results', []) if r.get('comprehension_pct', 0) >= 80),
        None
    )
    if rec_grade:
        row = next((r for r in payload.get('results', []) if r.get('grade') == rec_grade), None)
        for word in (row or {}).get('unknown_word_list', [])[:80]:
            items.append((analysis_id, 'difficult_word', word, 'Approve glossary/support word', rec_grade))
    for word in payload.get('proper_nouns', [])[:80]:
        items.append((analysis_id, 'proper_noun', word, 'Confirm as proper noun', rec_grade))
    for g in (suitability.get('glossary') or [])[:80]:
        word = g.get('word') or g.get('item') or ''
        suggestion = g.get('simple_meaning') or g.get('meaning') or g.get('definition') or 'Add simple definition'
        grade = g.get('level') or g.get('class_level') or rec_grade
        items.append((analysis_id, 'glossary', word, suggestion, grade))
    meaning = payload.get('meaning') or {}
    for f in (meaning.get('flagged') or [])[:60]:
        items.append((analysis_id, 'concept', f.get('item') or '', f"Review {f.get('severity', 'advanced')} concept", f.get('level') or rec_grade))
    seen, out = set(), []
    for item in items:
        key = (item[1], item[2])
        if item[2] and key not in seen:
            seen.add(key); out.append(item)
    return out

def _word_distribution_from_results(results):
    rows = []
    total = int(results[0].get('total_unique_book_words', 0)) if results else 0
    for r in results:
        words = r.get('new_at_grade_list') or []
        rows.append({
            'grade': r.get('grade'),
            'word_count': int(r.get('new_at_grade', len(words)) or 0),
            'word_pct': r.get('new_at_grade_pct', round(len(words) / total * 100, 1) if total else 0),
            'words': words,
        })
    if results:
        last = results[-1]
        words = last.get('unknown_word_list') or last.get('new_word_list') or []
        rows.append({
            'grade': None,
            'label': 'Not found in any class',
            'word_count': int(last.get('new_words', last.get('not_found_words', len(words))) or 0),
            'word_pct': last.get('new_pct', last.get('not_found_pct', round(len(words) / total * 100, 1) if total else 0)),
            'words': words,
        })
    return rows

@app.route('/api/review/<int:analysis_id>')
def review_items(analysis_id):
    conn = get_db()
    rows = conn.execute(
        'SELECT * FROM review_items WHERE analysis_id = ? ORDER BY status, item_type, id',
        (analysis_id,)
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/review/from_analysis/<int:analysis_id>', methods=['POST'])
def review_from_analysis(analysis_id):
    payload = _load_analysis_payload(analysis_id)
    if not payload:
        return jsonify({'error': 'Analysis not found'}), 404
    candidates = _review_candidates_from_analysis(payload)
    now = datetime.datetime.now().isoformat()
    conn = get_db()
    conn.execute('DELETE FROM review_items WHERE analysis_id = ? AND status = ?', (analysis_id, 'pending'))
    conn.executemany('''
        INSERT INTO review_items
          (analysis_id, item_type, item_text, suggestion, grade, status, created_at, updated_at)
        VALUES (?,?,?,?,?,'pending',?,?)
    ''', [(a, t, text, sug, grade, now, now) for a, t, text, sug, grade in candidates])
    conn.commit()
    rows = conn.execute(
        'SELECT * FROM review_items WHERE analysis_id = ? ORDER BY status, item_type, id',
        (analysis_id,)
    ).fetchall()
    conn.close()
    return jsonify({'ok': True, 'count': len(rows), 'items': [dict(r) for r in rows]})

@app.route('/api/review/item/<int:item_id>', methods=['POST'])
def review_update_item(item_id):
    data = request.json or {}
    status = data.get('status')
    if status not in {'pending', 'approved', 'rejected'}:
        return jsonify({'error': 'Invalid status'}), 400
    conn = get_db()
    conn.execute(
        'UPDATE review_items SET status = ?, notes = ?, updated_at = ? WHERE id = ?',
        (status, data.get('notes', ''), datetime.datetime.now().isoformat(), item_id)
    )
    conn.commit()
    row = conn.execute('SELECT * FROM review_items WHERE id = ?', (item_id,)).fetchone()
    conn.close()
    return jsonify({'ok': True, 'item': dict(row) if row else None})

@app.route('/api/glossary/<int:analysis_id>')
def glossary_items(analysis_id):
    conn = get_db()
    rows = conn.execute(
        'SELECT * FROM book_glossary WHERE analysis_id = ? ORDER BY status, grade, word',
        (analysis_id,)
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/glossary/from_analysis/<int:analysis_id>', methods=['POST'])
def glossary_from_analysis(analysis_id):
    payload = _load_analysis_payload(analysis_id)
    if not payload:
        return jsonify({'error': 'Analysis not found'}), 404
    suitability = payload.get('suitability') or {}
    rows = []
    for g in (suitability.get('glossary') or [])[:120]:
        word = g.get('word') or g.get('item')
        if word:
            rows.append((analysis_id, word, g.get('simple_meaning') or g.get('meaning') or g.get('definition') or '', g.get('level') or g.get('class_level') or suitability.get('recommended_grade'), 'analysis'))
    if not rows:
        meaning = payload.get('meaning') or {}
        for f in (meaning.get('flagged') or [])[:120]:
            word = f.get('item')
            if word:
                rows.append((analysis_id, word, f.get('concept') or 'Teacher explanation needed', f.get('level'), 'meaning'))
    now = datetime.datetime.now().isoformat()
    conn = get_db()
    conn.execute('DELETE FROM book_glossary WHERE analysis_id = ? AND status = ?', (analysis_id, 'draft'))
    conn.executemany('''
        INSERT INTO book_glossary
          (analysis_id, word, definition, grade, source, status, created_at, updated_at)
        VALUES (?,?,?,?,?,'draft',?,?)
    ''', [(a, w, definition, grade, source, now, now) for a, w, definition, grade, source in rows])
    conn.commit()
    saved = conn.execute(
        'SELECT * FROM book_glossary WHERE analysis_id = ? ORDER BY status, grade, word',
        (analysis_id,)
    ).fetchall()
    conn.close()
    return jsonify({'ok': True, 'count': len(saved), 'items': [dict(r) for r in saved]})

@app.route('/api/glossary/item/<int:item_id>', methods=['POST'])
def glossary_update_item(item_id):
    data = request.json or {}
    conn = get_db()
    conn.execute('''
        UPDATE book_glossary
        SET word = ?, definition = ?, grade = ?, status = ?, updated_at = ?
        WHERE id = ?
    ''', (
        data.get('word', '').strip(),
        data.get('definition', '').strip(),
        data.get('grade') or None,
        data.get('status') or 'draft',
        datetime.datetime.now().isoformat(),
        item_id,
    ))
    conn.commit()
    row = conn.execute('SELECT * FROM book_glossary WHERE id = ?', (item_id,)).fetchone()
    conn.close()
    return jsonify({'ok': True, 'item': dict(row) if row else None})

@app.route('/api/admin/dashboard')
def admin_dashboard():
    conn = get_db()
    grades = [dict(r) for r in conn.execute('SELECT * FROM grade_meta ORDER BY grade').fetchall()]
    files = [dict(r) for r in conn.execute('SELECT grade, filename, word_count, processed_at, source FROM grade_files ORDER BY grade, filename').fetchall()]
    totals = dict(conn.execute('''
        SELECT
          (SELECT COUNT(*) FROM grade_files) AS grade_files,
          (SELECT COUNT(*) FROM analyses) AS analyses,
          (SELECT COUNT(*) FROM grade_words) AS grade_words,
          (SELECT COUNT(*) FROM review_items WHERE status = 'pending') AS pending_reviews,
          (SELECT COUNT(*) FROM book_glossary) AS glossary_items,
          (SELECT COUNT(*) FROM reading_attempts) AS reading_attempts
    ''').fetchone())
    conn.close()
    missing = [g for g in range(1, 13) if not any(int(row.get('grade', 0)) == g for row in grades)]
    cache_dir = os.path.join('data', 'cache')
    cache_files = [p for p in os.listdir(cache_dir)] if os.path.isdir(cache_dir) else []
    cfg = _fw.load_config() if '_fw' in globals() else {}
    return jsonify({
        'totals': totals,
        'grades': grades,
        'files': files[:300],
        'missing_grades': missing,
        'db_size_mb': round(os.path.getsize(DB_PATH) / (1024 * 1024), 2) if os.path.exists(DB_PATH) else 0,
        'cache_files': len(cache_files),
        'watch_folder': cfg.get('watch_folder', ''),
        'watcher_status': dict(_fw.WATCHER_STATUS) if '_fw' in globals() else {},
    })

READING_PASSAGES = {
    1: [
        'அம்மா வீட்டில் இருக்கிறார். நான் புத்தகம் படிக்கிறேன். மரம் அருகில் பறவை உள்ளது.',
        'பூனை பால் குடிக்கிறது. குழந்தை படம் பார்க்கிறது. அப்பா கதை சொல்கிறார்.',
    ],
    2: [
        'குழந்தைகள் பள்ளிக்கு சென்றனர். ஆசிரியர் நல்ல கதையைப் படித்தார். அனைவரும் கவனமாக கேட்டனர்.',
        'மழை பெய்தது. மரங்கள் பசுமையாக இருந்தன. மாணவர்கள் மகிழ்ச்சியுடன் பாடம் படித்தனர்.',
    ],
    3: [
        'எங்கள் ஊரில் அழகான பூங்கா உள்ளது. அங்கு குழந்தைகள் மாலை நேரத்தில் விளையாடுகிறார்கள். மரங்களும் மலர்களும் அந்த இடத்தை அழகாக்குகின்றன.',
        'காலை நேரத்தில் சூரியன் உதிக்கிறது. பறவைகள் இனிமையாக பாடுகின்றன. மாணவர்கள் பள்ளிக்குச் செல்லத் தயாராகிறார்கள்.',
    ],
    4: [
        'நீர் நம் வாழ்விற்கு மிகவும் அவசியமானது. நாம் குடிக்கவும் சமைக்கவும் விவசாயம் செய்யவும் நீரைப் பயன்படுத்துகிறோம். ஆகவே நீரை வீணாக்காமல் பாதுகாக்க வேண்டும்.',
        'நூலகம் அறிவை வளர்க்கும் சிறந்த இடம். அங்கு பல வகையான புத்தகங்கள் உள்ளன. நல்ல புத்தகங்களைப் படிப்பதால் சிந்தனை திறன் வளரும்.',
    ],
    5: [
        'சுற்றுச்சூழலைப் பாதுகாப்பது ஒவ்வொருவரின் கடமை. மரங்களை நடுதல், நீரைச் சேமித்தல், குப்பையை சரியான இடத்தில் போடுதல் போன்ற பழக்கங்கள் நம் ஊரை தூய்மையாக வைத்திருக்கும்.',
    ],
    6: [
        'தமிழ் மொழி பழமையான செம்மொழிகளில் ஒன்றாகும். அதன் இலக்கியங்கள் மனித வாழ்க்கை, இயற்கை, அறம், அறிவு ஆகியவற்றைப் பற்றி அழகாக எடுத்துரைக்கின்றன.',
    ],
    7: [
        'அறிவியல் சிந்தனை மனிதனுக்கு காரணத்தை ஆராயும் திறனை அளிக்கிறது. ஒரு நிகழ்வு ஏன் நடக்கிறது என்பதை கேள்வி கேட்டு ஆராயும்போது புதிய கண்டுபிடிப்புகள் உருவாகின்றன.',
    ],
    8: [
        'சமூகத்தில் ஒற்றுமை நிலைக்க வேண்டுமெனில் அனைவரும் ஒருவரை ஒருவர் மதிக்க வேண்டும். மொழி, மதம், பழக்கம் ஆகிய வேறுபாடுகள் இருந்தாலும் மனித நேயம் பொதுவான மதிப்பாக இருக்க வேண்டும்.',
    ],
    9: [
        'வரலாற்றைப் படிப்பது கடந்த கால நிகழ்வுகளை அறிதலுக்கு மட்டுமல்ல; தற்போதைய சமூக மாற்றங்களைப் புரிந்துகொள்வதற்கும் உதவுகிறது. மக்கள் எடுத்த முடிவுகள் எதிர்காலத்தை எவ்வாறு பாதித்தன என்பதையும் அது காட்டுகிறது.',
    ],
    10: [
        'தொழில்நுட்ப வளர்ச்சி கல்வி முறையில் பல மாற்றங்களை ஏற்படுத்தியுள்ளது. இணையம் மூலம் மாணவர்கள் பல்வேறு அறிவு வளங்களை அணுக முடிகிறது. ஆனால் தகவலை சிந்தித்து தேர்ந்தெடுக்கும் திறனும் அவசியம்.',
    ],
    11: [
        'இலக்கியப் படைப்புகள் சமூகத்தின் உணர்வுகளையும் முரண்பாடுகளையும் வெளிப்படுத்தும் ஆற்றல் கொண்டவை. ஒரு சிறந்த படைப்பு வாசகரை சிந்திக்கவும் தன் அனுபவத்தை மறுபரிசீலனை செய்யவும் தூண்டும்.',
    ],
    12: [
        'மனித முன்னேற்றம் அறிவு, பொறுப்பு, கருணை ஆகிய மூன்றின் சமநிலையால் நிலைபெறும். அறிவியல் புதிய வாய்ப்புகளைத் திறந்தாலும், அவற்றை அறநெறியுடன் பயன்படுத்தும் சமூகப் பொறுப்பு அவசியமானது.',
    ],
}

def _reading_passage_for_grade(grade):
    grade = max(1, min(12, int(grade or 1)))
    items = READING_PASSAGES.get(grade) or READING_PASSAGES[12]
    return items[int(time.time()) % len(items)]

@app.route('/api/reading/passage')
def reading_passage():
    grade = int(request.args.get('grade') or 1)
    text = _reading_passage_for_grade(grade)
    return jsonify({'grade': grade, 'text': text, 'word_count': len(_reading_score.tamil_words(text))})

@app.route('/api/reading/asr_status')
def reading_asr_status():
    return jsonify({
        'configured': bool(os.environ.get('TAMIL_READING_ASR_CMD') or (os.environ.get('WHISPER_CPP_BIN') and os.environ.get('WHISPER_CPP_MODEL'))),
        'engine': 'custom' if os.environ.get('TAMIL_READING_ASR_CMD') else 'whisper.cpp',
        'whisper_bin': os.environ.get('WHISPER_CPP_BIN', ''),
        'model': os.path.basename(os.environ.get('WHISPER_CPP_MODEL', '')),
    })

@app.route('/api/reading/submit', methods=['POST'])
def reading_submit():
    grade = int(request.form.get('grade') or 1)
    expected_text = request.form.get('passage_text', '')
    student_name = request.form.get('student_name', '').strip()
    strictness = request.form.get('strictness', 'gentle')
    manual_transcript = request.form.get('manual_transcript', '').strip()

    if not expected_text.strip():
        return jsonify({'error': 'Missing passage text.'}), 400

    audio_path = ''
    transcript = manual_transcript
    engine = 'manual'
    try:
        if not transcript:
            if 'audio' not in request.files:
                return jsonify({'error': 'No audio recording received.'}), 400
            audio = request.files['audio']
            os.makedirs('uploads/reading', exist_ok=True)
            safe = secure_filename(audio.filename or 'reading.webm') or 'reading.webm'
            audio_path = os.path.join('uploads', 'reading', f'{uuid.uuid4().hex}_{safe}')
            audio.save(audio_path)
            asr = _reading_asr.transcribe(audio_path, 'ta')
            transcript = asr.get('transcript', '')
            engine = asr.get('engine', 'asr')
            if audio_path and os.path.exists(audio_path):
                os.unlink(audio_path)
                audio_path = ''

        score = _reading_score.score_reading(
            expected_text,
            transcript,
            stem_fn=get_stem,
            strictness=strictness,
        )
        conn = get_db()
        conn.execute('''
            INSERT INTO reading_attempts
              (student_name, grade, passage_text, transcript, engine, strictness, score_json, audio_path, created_at)
            VALUES (?,?,?,?,?,?,?,?,?)
        ''', (
            student_name, grade, expected_text, transcript, engine, strictness,
            json.dumps(score, ensure_ascii=False), '', datetime.datetime.now().isoformat()
        ))
        conn.commit()
        attempt_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
        conn.close()
        if audio_path and os.path.exists(audio_path):
            os.unlink(audio_path)
        return jsonify({'ok': True, 'attempt_id': attempt_id, 'engine': engine, 'transcript': transcript, 'score': score})
    except _reading_asr.ASRNotConfigured as e:
        if audio_path and os.path.exists(audio_path):
            os.unlink(audio_path)
        return jsonify({'error': str(e), 'asr_not_configured': True}), 503
    except Exception as e:
        logging.getLogger('app').exception('Reading assessment failed')
        if audio_path and os.path.exists(audio_path):
            os.unlink(audio_path)
        return jsonify({'error': str(e)}), 500

@app.route('/api/reading/history')
def reading_history():
    grade = request.args.get('grade', '').strip()
    student = request.args.get('student', '').strip()
    params = []
    where = []
    if grade:
        where.append('grade = ?')
        params.append(int(grade))
    if student:
        where.append('LOWER(student_name) LIKE ?')
        params.append(f'%{student.lower()}%')
    clause = ('WHERE ' + ' AND '.join(where)) if where else ''
    conn = get_db()
    rows = conn.execute(f'''
        SELECT id, student_name, grade, passage_text, transcript, engine, strictness, score_json, created_at
        FROM reading_attempts {clause} ORDER BY id DESC LIMIT 200
    ''', params).fetchall()
    conn.close()
    out = []
    for r in rows:
        d = dict(r)
        score = json.loads(d.pop('score_json') or '{}')
        d['final_mark'] = score.get('final_mark')
        d['reading_accuracy'] = score.get('reading_accuracy')
        d['pronunciation_confidence'] = score.get('pronunciation_confidence')
        d['counts'] = score.get('counts', {})
        d['practice_words'] = score.get('practice_words', [])
        out.append(d)
    return jsonify(out)

@app.route('/api/reading/clear', methods=['POST'])
def reading_clear():
    conn = get_db()
    conn.execute('DELETE FROM reading_attempts')
    conn.commit()
    conn.close()
    # Defensive cleanup for old temporary recordings from earlier builds.
    reading_dir = os.path.join('uploads', 'reading')
    if os.path.isdir(reading_dir):
        for name in os.listdir(reading_dir):
            path = os.path.join(reading_dir, name)
            try:
                if os.path.isfile(path):
                    os.unlink(path)
            except Exception:
                pass
    return jsonify({'ok': True})

@app.route('/api/tamil_features/analyze', methods=['POST'])
def tamil_features_analyze():
    data = request.json or {}
    text = data.get('text', '')
    if not text.strip():
        return jsonify({'error': 'Tamil text is required.'}), 400
    try:
        return jsonify(_tamil_features.analyze(text, stem_fn=get_stem))
    except Exception as e:
        logging.getLogger('app').exception('Tamil feature analysis failed')
        return jsonify({'error': str(e)}), 500

@app.route('/api/history/<int:analysis_id>', methods=['DELETE'])
def delete_analysis(analysis_id):
    conn = get_db()
    conn.execute('DELETE FROM analyses WHERE id = ?', (analysis_id,))
    conn.commit(); conn.close()
    return jsonify({'ok': True})

# ── Report generation ─────────────────────────────────────────────────────────

@app.route('/api/report/<int:analysis_id>')
def generate_report(analysis_id):
    conn = get_db()
    row = conn.execute('SELECT * FROM analyses WHERE id = ?', (analysis_id,)).fetchone()
    conn.close()
    if not row: return jsonify({'error': 'Not found'}), 404

    row = dict(row)
    results      = json.loads(row['results_json'])
    distribution = _word_distribution_from_results(results)
    proper_nouns = json.loads(row.get('proper_nouns') or '[]')
    sent_data    = json.loads(row.get('sentence_json') or '{}')
    book_name    = row['book_name']
    tss          = sent_data.get('target', {})
    meaning      = json.loads(row.get('meaning_json') or 'null') if 'meaning_json' in row.keys() else None

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        Table, TableStyle, HRFlowable,
                                        KeepTogether, PageBreak)
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # ── Tamil font registration ───────────────────────────────────────────
        # Priority order:
        #   1. Bundled fonts/ folder (ships with the app — most reliable)
        #   2. Common system font paths (Linux / Windows / macOS)
        #   3. Helvetica fallback (Latin only — Tamil will not render)

        _base = os.path.dirname(os.path.abspath(__file__))
        _font_candidates = [
            # Bundled with app
            os.path.join(_base, 'fonts', 'NotoSansTamil-Regular.ttf'),
            os.path.join(_base, 'fonts', 'FreeSerif.ttf'),
            # Linux system paths
            '/usr/share/fonts/truetype/freefont/FreeSerif.ttf',
            '/usr/share/fonts/truetype/noto/NotoSansTamil-Regular.ttf',
            '/usr/share/fonts/truetype/lohit-tamil/Lohit-Tamil.ttf',
            '/usr/share/fonts/noto/NotoSansTamil-Regular.ttf',
            # Windows paths
            r'C:\Windows\Fonts\nirmala.ttf',
            r'C:\Windows\Fonts\NotoSansTamil-Regular.ttf',
            # macOS
            '/System/Library/Fonts/Supplemental/NotoSansTamil-Regular.ttf',
            '/Library/Fonts/NotoSansTamil-Regular.ttf',
        ]
        _bold_candidates = [
            os.path.join(_base, 'fonts', 'FreeSerifBold.ttf'),
            '/usr/share/fonts/truetype/freefont/FreeSerifBold.ttf',
        ]

        TAMIL_FONT      = 'Helvetica'
        TAMIL_FONT_BOLD = 'Helvetica-Bold'

        for fp in _font_candidates:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont('TamilFont', fp))
                    TAMIL_FONT = 'TamilFont'
                    break
                except: pass

        if TAMIL_FONT != 'Helvetica':
            for fp in _bold_candidates:
                if os.path.exists(fp):
                    try:
                        pdfmetrics.registerFont(TTFont('TamilFontBold', fp))
                        TAMIL_FONT_BOLD = 'TamilFontBold'
                        break
                    except: pass

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)

        styles = getSampleStyleSheet()
        W = A4[0] - 4*cm

        # Custom styles
        def S(name, **kw):
            return ParagraphStyle(name, fontName=kw.pop('fontName', 'Helvetica'),
                                  fontSize=kw.pop('fontSize', 10),
                                  leading=kw.pop('leading', 14),
                                  textColor=kw.pop('textColor', colors.black),
                                  alignment=kw.pop('alignment', TA_LEFT), **kw)

        sTitle    = S('sTitle',    fontName='Helvetica-Bold', fontSize=18, leading=24, alignment=TA_CENTER, textColor=colors.HexColor('#1a3a5c'), spaceAfter=4)
        sSub      = S('sSub',      fontName=TAMIL_FONT,       fontSize=10, leading=16, alignment=TA_CENTER, textColor=colors.HexColor('#5c574f'), spaceAfter=12)
        sH1       = S('sH1',       fontName='Helvetica-Bold', fontSize=13, leading=18, textColor=colors.HexColor('#1D4E89'), spaceBefore=14, spaceAfter=4)
        sH2       = S('sH2',       fontName=TAMIL_FONT_BOLD,  fontSize=11, leading=16, textColor=colors.HexColor('#2e6da4'), spaceBefore=10, spaceAfter=3)
        sBody     = S('sBody',     fontName=TAMIL_FONT,       fontSize=9,  leading=14, spaceAfter=3)
        sNote     = S('sNote',     fontName=TAMIL_FONT,       fontSize=8,  leading=12, textColor=colors.HexColor('#777777'), spaceAfter=6)
        sTamil    = S('sTamil',    fontName=TAMIL_FONT,       fontSize=10, leading=15)
        sWordList = S('sWordList', fontName=TAMIL_FONT,       fontSize=10, leading=16, spaceAfter=6)

        # Color palette
        COL_HEADER = colors.HexColor('#1D4E89')
        COL_GREEN  = colors.HexColor('#D4EDDA')
        COL_AMBER  = colors.HexColor('#FFF3CD')
        COL_RED    = colors.HexColor('#F8D7DA')
        COL_BLUE   = colors.HexColor('#D6EAF8')
        COL_GRAY   = colors.HexColor('#F0EDE8')
        COL_PURPLE = colors.HexColor('#EDE0FB')

        def verdict_color(pct):
            if pct >= 80: return COL_GREEN
            if pct >= 60: return COL_AMBER
            return COL_RED

        story = []

        # ── Cover ──────────────────────────────────────────────────────────────
        story.append(Spacer(1, 1*cm))
        story.append(Paragraph('Tamil Book Readability Report', sTitle))
        story.append(Paragraph(book_name, sSub))
        story.append(Paragraph(f'Generated: {row["analyzed_at"][:19]}', sSub))
        story.append(HRFlowable(width=W, thickness=1, color=COL_HEADER, spaceAfter=12))

        # ── Overview stats ─────────────────────────────────────────────────────
        story.append(Paragraph('Overview', sH1))
        first_readable = next((r for r in results if r['comprehension_pct'] >= 80), None)
        total_stems = row['unique_stems'] or row['unique_words']

        ov_data = [
            ['Total words in book', f"{row['total_words']:,}"],
            ['Unique Tamil stems (after morphological analysis)', f"{total_stems:,}"],
            ['Proper nouns (counted as known)', str(len(proper_nouns))],
            ['First readable from', f"Standard {first_readable['grade']}" if first_readable else 'Beyond Std 12'],
            ['Target book — average words per sentence', str(tss.get('avg', '—'))],
            ['Target book — max words in one sentence', str(tss.get('max', '—'))],
            ['Target book — total sentences analysed', f"{tss.get('total_sentences', 0):,}"],
        ]
        ov_table = Table(ov_data, colWidths=[W*0.65, W*0.35])
        ov_table.setStyle(TableStyle([
            ('FONTNAME',    (0,0),(-1,-1), 'Helvetica'),
            ('FONTSIZE',    (0,0),(-1,-1), 9),
            ('FONTNAME',    (0,0),(0,-1),  'Helvetica-Bold'),
            ('BACKGROUND',  (0,0),(-1,-1), COL_GRAY),
            ('BACKGROUND',  (0,0),(-1, 0), COL_BLUE),
            ('ROWBACKGROUNDS',(0,0),(-1,-1),[colors.white, COL_GRAY]),
            ('GRID',        (0,0),(-1,-1), 0.3, colors.HexColor('#cccccc')),
            ('TOPPADDING',  (0,0),(-1,-1), 5),
            ('BOTTOMPADDING',(0,0),(-1,-1), 5),
            ('LEFTPADDING', (0,0),(-1,-1), 8),
        ]))
        story.append(ov_table)
        story.append(Spacer(1, 0.4*cm))

        # ── Main readability table ─────────────────────────────────────────────
        story.append(Paragraph('Section 1 — Readability by Standard', sH1))
        story.append(Paragraph(
            'Each row is cumulative: "Std 1–5" means a student who has completed Standards 1 through 5. '
            '"Known words" = words from the book the student already knows (Std 1 through N combined). '
            '"New words for student" = words in the book the student has not yet learned.',
            sNote))

        hdr = [
            'Standard',
            'Total unique\nwords (book)',
            'Known words\n(Std 1–N)',
            '% known',
            'New words\n(new to student)',
            '% new',
            'Verdict',
            'Grade max\nsentence',
            'Sentences\nover max',
        ]
        tbl_data = [hdr]
        for r in results:
            pct = r['known_pct']
            verdict = 'Easy' if pct>=90 else 'Readable' if pct>=80 else 'Challenging' if pct>=60 else 'Very Hard'
            g = r['grade']
            tbl_data.append([
                f'Std 1–{g}',
                f"{r['total_unique_book_words']:,}",
                f"{r['known_words']:,}",
                f"{r['known_pct']}%",
                f"{r['new_words']:,}",
                f"{r['new_pct']}%",
                verdict,
                str(r.get('grade_sent_max', '—')),
                str(r.get('target_sentences_over_max', '—')),
            ])

        col_w = [W*x for x in [0.12, 0.13, 0.12, 0.08, 0.13, 0.08, 0.13, 0.11, 0.10]]
        main_tbl = Table(tbl_data, colWidths=col_w, repeatRows=1)
        tbl_style = [
            ('FONTNAME',     (0,0),(-1, 0), 'Helvetica-Bold'),
            ('FONTNAME',     (0,1),(-1,-1), TAMIL_FONT),
            ('FONTSIZE',     (0,0),(-1,-1), 8),
            ('BACKGROUND',   (0,0),(-1, 0), COL_HEADER),
            ('TEXTCOLOR',    (0,0),(-1, 0), colors.white),
            ('ALIGN',        (0,0),(-1,-1), 'CENTER'),
            ('VALIGN',       (0,0),(-1,-1), 'MIDDLE'),
            ('GRID',         (0,0),(-1,-1), 0.3, colors.HexColor('#bbbbbb')),
            ('TOPPADDING',   (0,0),(-1,-1), 4),
            ('BOTTOMPADDING',(0,0),(-1,-1), 4),
            ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white, COL_GRAY]),
        ]
        # Color-code verdict column and comprehension column
        for i, r in enumerate(results, 1):
            pct = r['comprehension_pct']
            c = verdict_color(pct)
            tbl_style.append(('BACKGROUND', (6,i),(6,i), c))
            tbl_style.append(('BACKGROUND', (3,i),(3,i), c))
        main_tbl.setStyle(TableStyle(tbl_style))
        story.append(main_tbl)
        story.append(Spacer(1, 0.4*cm))
        story.append(Paragraph(
            '"Known words" = words from this book a student at Std 1–N already knows. '
            '"New words for student" = words in this book the student has not yet learned. '
            '"Sentences over max" = sentences longer than the longest sentence in that standard\'s textbook.',
            sNote))

        story.append(Paragraph('Word Distribution by Class', sH1))
        story.append(Paragraph(
            'Each unique book word is assigned to the first class where it appears in the loaded textbook database. '
            'The final row lists words not found in any class.',
            sNote))
        dist_data = [['Class', 'Words', '% of book vocabulary']]
        for drow in distribution:
            label = drow.get('label') or f"Std {drow.get('grade')}"
            dist_data.append([label, f"{drow.get('word_count', 0):,}", f"{drow.get('word_pct', 0)}%"])
        dist_tbl = Table(dist_data, colWidths=[W*0.40, W*0.25, W*0.25])
        dist_tbl.setStyle(TableStyle([
            ('FONTNAME',     (0,0),(-1, 0), 'Helvetica-Bold'),
            ('FONTNAME',     (0,1),(-1,-1), TAMIL_FONT),
            ('FONTSIZE',     (0,0),(-1,-1), 9),
            ('BACKGROUND',   (0,0),(-1, 0), COL_HEADER),
            ('TEXTCOLOR',    (0,0),(-1, 0), colors.white),
            ('ALIGN',        (1,1),(-1,-1), 'CENTER'),
            ('GRID',         (0,0),(-1,-1), 0.3, colors.HexColor('#bbbbbb')),
            ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white, COL_GRAY]),
            ('TOPPADDING',   (0,0),(-1,-1), 5),
            ('BOTTOMPADDING',(0,0),(-1,-1), 5),
        ]))
        story.append(dist_tbl)

        # ── Sentence complexity section ────────────────────────────────────────
        story.append(PageBreak())
        story.append(Paragraph('Section 2 — Sentence Complexity Analysis', sH1))
        story.append(Paragraph(
            'Compares sentence length distribution of the target book against each school standard\'s textbook. '
            'Long sentences are a major contributor to readability difficulty beyond vocabulary.',
            sNote))

        sent_hdr = ['Standard', 'Grade book\nmax sentence', 'Grade book\navg sentence',
                    'Target book\navg sentence', 'Sentences in\ntarget > grade max',
                    '% sentences\nover grade max']
        sent_data_rows = [sent_hdr]
        target_avg = tss.get('avg', 0)
        for r in results:
            gmax = r.get('grade_sent_max', 0)
            over = r.get('target_sentences_over_max', 0)
            over_pct = r.get('target_pct_over_max', 0)
            difficulty = ''
            if gmax > 0:
                if over_pct > 50: difficulty = 'High'
                elif over_pct > 20: difficulty = 'Medium'
                else: difficulty = 'Low'
            sent_data_rows.append([
                f"Std {r['grade']}",
                str(gmax) if gmax else '—',
                str(r.get('grade_sent_avg','—')),
                str(target_avg),
                str(over),
                f"{over_pct}%  [{difficulty}]",
            ])

        s_col_w = [W*x for x in [0.1, 0.16, 0.16, 0.16, 0.22, 0.20]]
        sent_tbl = Table(sent_data_rows, colWidths=s_col_w, repeatRows=1)
        sent_style = [
            ('FONTNAME',     (0,0),(-1, 0), 'Helvetica-Bold'),
            ('FONTNAME',     (0,1),(-1,-1), TAMIL_FONT),
            ('FONTSIZE',     (0,0),(-1,-1), 8.5),
            ('BACKGROUND',   (0,0),(-1, 0), COL_HEADER),
            ('TEXTCOLOR',    (0,0),(-1, 0), colors.white),
            ('ALIGN',        (1,0),(-1,-1), 'CENTER'),
            ('ALIGN',        (0,0),(0,-1),  'LEFT'),
            ('VALIGN',       (0,0),(-1,-1), 'MIDDLE'),
            ('GRID',         (0,0),(-1,-1), 0.3, colors.HexColor('#bbbbbb')),
            ('TOPPADDING',   (0,0),(-1,-1), 5),
            ('BOTTOMPADDING',(0,0),(-1,-1), 5),
            ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white, COL_GRAY]),
        ]
        sent_tbl.setStyle(TableStyle(sent_style))
        story.append(sent_tbl)
        story.append(Spacer(1, 0.3*cm))
        story.append(Paragraph(
            f'Target book sentence stats — '
            f'Average: {tss.get("avg","—")} words/sentence | '
            f'Max: {tss.get("max","—")} words | '
            f'Median: {tss.get("median","—")} words | '
            f'Total sentences: {tss.get("total_sentences",0):,}',
            sNote))

        # ── Per-grade new word lists ───────────────────────────────────────────
        story.append(PageBreak())
        story.append(Paragraph('Section 3 — New Words Introduced per Standard', sH1))
        story.append(Paragraph(
            'For each standard, the words in this book that become "newly known" after completing that standard. '
            'These represent the vocabulary a student gains at each level that directly helps reading this book.',
            sNote))

        for r in results:
            nw = r.get('new_at_grade_list', r.get('new_word_list', []))
            if not nw: continue
            pct = r.get('new_at_grade_pct', r.get('new_words_pct', 0))
            story.append(KeepTogether([
                Paragraph(f"Standard 1\u2013{r['grade']} \u2014 {len(nw)} words introduced at Std {r['grade']} ({pct}% of book vocabulary)", sH2),
            ]))
            # Display as a compact word grid using a table
            COLS = 6
            rows_of_words = []
            row_w = []
            for i in range(0, len(nw), COLS):
                chunk = nw[i:i+COLS]
                while len(chunk) < COLS: chunk.append('')
                rows_of_words.append([Paragraph(w, sTamil) for w in chunk])
            if rows_of_words:
                col_w_words = [W/COLS]*COLS
                wt = Table(rows_of_words, colWidths=col_w_words)
                wt.setStyle(TableStyle([
                    ('FONTNAME',    (0,0),(-1,-1), TAMIL_FONT),
                    ('FONTSIZE',    (0,0),(-1,-1), 9),
                    ('TOPPADDING',  (0,0),(-1,-1), 2),
                    ('BOTTOMPADDING',(0,0),(-1,-1), 2),
                    ('LEFTPADDING', (0,0),(-1,-1), 4),
                    ('ROWBACKGROUNDS',(0,0),(-1,-1),[colors.white, COL_GRAY]),
                    ('GRID',        (0,0),(-1,-1), 0.2, colors.HexColor('#dddddd')),
                ]))
                story.append(wt)
            story.append(Spacer(1, 0.2*cm))

        # ── Unknown words per grade ────────────────────────────────────────────
        story.append(PageBreak())
        story.append(Paragraph('Section 4 — New Words for Student per Standard', sH1))
        story.append(Paragraph(
            'Words in this book that are new to a student at each level — '
            'not yet encountered in their studies up to that standard.',
            sNote))

        for r in results:
            uw = r.get('new_word_list', r.get('unknown_word_list', []))
            story.append(KeepTogether([
                Paragraph(
                    f"Standard 1\u2013{r['grade']} \u2014 "
                    f"{r.get('new_words', r.get('unknown_words', 0))} new words for student "
                    f"({r.get('new_pct', round(100-r.get('known_pct',r.get('comprehension_pct',100)),1))}% of book vocabulary)",
                    sH2),
            ]))
            if not uw:
                story.append(Paragraph('No new words — this student already knows every word in this book.', sNote))
                continue
            COLS = 6
            rows_uw = []
            for i in range(0, min(len(uw), 300), COLS):
                chunk = uw[i:i+COLS]
                while len(chunk) < COLS: chunk.append('')
                rows_uw.append([Paragraph(w, sTamil) for w in chunk])
            if rows_uw:
                uwt = Table(rows_uw, colWidths=[W/COLS]*COLS)
                uwt.setStyle(TableStyle([
                    ('FONTNAME',    (0,0),(-1,-1), TAMIL_FONT),
                    ('FONTSIZE',    (0,0),(-1,-1), 9),
                    ('TOPPADDING',  (0,0),(-1,-1), 2),
                    ('BOTTOMPADDING',(0,0),(-1,-1), 2),
                    ('LEFTPADDING', (0,0),(-1,-1), 4),
                    ('ROWBACKGROUNDS',(0,0),(-1,-1),[colors.white, COL_RED]),
                    ('GRID',        (0,0),(-1,-1), 0.2, colors.HexColor('#dddddd')),
                ]))
                story.append(uwt)
            if len(uw) > 300:
                story.append(Paragraph(f'… and {len(uw)-300} more (see Excel export for full list)', sNote))
            story.append(Spacer(1, 0.2*cm))

        # ── Proper nouns ───────────────────────────────────────────────────────
        if proper_nouns:
            story.append(PageBreak())
            story.append(Paragraph('Section 5 — Proper Nouns (Counted as Known)', sH1))
            story.append(Paragraph(
                'These words were identified as names, places, deities, or foreign proper nouns. '
                'They are counted as known at all grade levels since students can learn them from context.',
                sNote))
            COLS = 5
            pn_rows = []
            for i in range(0, len(proper_nouns), COLS):
                chunk = proper_nouns[i:i+COLS]
                while len(chunk) < COLS: chunk.append('')
                pn_rows.append([Paragraph(w, sTamil) for w in chunk])
            if pn_rows:
                pnt = Table(pn_rows, colWidths=[W/COLS]*COLS)
                pnt.setStyle(TableStyle([
                    ('FONTNAME',    (0,0),(-1,-1), TAMIL_FONT),
                    ('FONTSIZE',    (0,0),(-1,-1), 9),
                    ('TOPPADDING',  (0,0),(-1,-1), 3),
                    ('BOTTOMPADDING',(0,0),(-1,-1), 3),
                    ('LEFTPADDING', (0,0),(-1,-1), 4),
                    ('ROWBACKGROUNDS',(0,0),(-1,-1),[colors.white, COL_PURPLE]),
                    ('GRID',        (0,0),(-1,-1), 0.2, colors.HexColor('#dddddd')),
                ]))
                story.append(pnt)

        doc.build(story)
        buf.seek(0)
        safe_name = re.sub(r'[^\w]', '_', os.path.splitext(book_name)[0])
        return send_file(buf, mimetype='application/pdf',
                         as_attachment=True,
                         download_name=f'report_{safe_name}.pdf')

    except ImportError as e:
        return jsonify({'error': f'reportlab not installed: {e}'}), 500

@app.route('/api/export/<int:analysis_id>')
def export_excel(analysis_id):
    conn = get_db()
    row = conn.execute('SELECT * FROM analyses WHERE id = ?', (analysis_id,)).fetchone()
    conn.close()
    if not row: return jsonify({'error': 'Not found'}), 404

    row = dict(row)
    results      = json.loads(row['results_json'])
    distribution = _word_distribution_from_results(results)
    proper_nouns = json.loads(row.get('proper_nouns') or '[]')
    sent_data    = json.loads(row.get('sentence_json') or '{}')
    book_name    = row['book_name']
    tss          = sent_data.get('target', {})
    meaning      = json.loads(row.get('meaning_json') or 'null') if 'meaning_json' in row.keys() else None

    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        hf   = Font(bold=True, color='FFFFFF')
        hfill= PatternFill('solid', fgColor='1D4E89')
        gf   = PatternFill('solid', fgColor='D4EDDA')
        af   = PatternFill('solid', fgColor='FFF3CD')
        rf   = PatternFill('solid', fgColor='F8D7DA')
        thin = Border(left=Side(style='thin'), right=Side(style='thin'),
                      top=Side(style='thin'), bottom=Side(style='thin'))
        def hdr_row(ws, headers, row_num=1):
            for col, h in enumerate(headers, 1):
                c = ws.cell(row=row_num, column=col, value=h)
                c.font = hf; c.fill = hfill
                c.alignment = Alignment(horizontal='center', wrap_text=True)
                c.border = thin

        # Sheet 1: Full summary table
        ws1 = wb.active; ws1.title = 'Readability Summary'
        ws1['A1'] = f'Tamil Readability Report — {book_name}'
        ws1['A1'].font = Font(bold=True, size=14)
        ws1['A2'] = (f'Analyzed: {row["analyzed_at"][:19]}  |  '
                     f'Total words: {row["total_words"]:,}  |  '
                     f'Unique stems: {row.get("unique_stems","—")}  |  '
                     f'Proper nouns: {len(proper_nouns)}  |  '
                     f'Target avg sentence: {tss.get("avg","—")} words  |  '
                     f'Target max sentence: {tss.get("max","—")} words')
        ws1['A2'].font = Font(italic=True, color='555555', size=9)
        ws1.append([])
        headers1 = [
            'Standard (cumulative)',
            'Total unique words in book',
            'Known words (Std 1–N)',
            '% known',
            'New words (new to student)',
            '% new',
            'Verdict',
            'Grade max sentence',
            'Book sentences over grade max',
        ]
        hdr_row(ws1, headers1, 4)
        for r in results:
            pct = r['known_pct']
            verdict = 'Easy' if pct>=90 else 'Readable' if pct>=80 else 'Challenging' if pct>=60 else 'Very Hard'
            ws1.append([
                f"Std 1–{r['grade']}",
                r['total_unique_book_words'],
                r['known_words'],
                f"{r['known_pct']}%",
                r['new_words'],
                f"{r['new_pct']}%",
                verdict,
                r.get('grade_sent_max','—'),
                r.get('target_sentences_over_max','—'),
            ])
            rn = ws1.max_row
            fill = gf if pct>=80 else af if pct>=60 else rf
            for col in range(1, 10):
                c = ws1.cell(row=rn, column=col)
                c.fill = fill; c.border = thin
                c.alignment = Alignment(horizontal='center')
        for col, w in enumerate([16, 22, 20, 10, 22, 10, 14, 16, 22], 1):
            ws1.column_dimensions[get_column_letter(col)].width = w

        # Sheet 2: Word distribution by class
        ws_dist = wb.create_sheet('Word Distribution')
        hdr_row(ws_dist, ['Class', 'Word count', '% of book vocabulary', 'Word'])
        for drow in distribution:
            label = drow.get('label') or f"Std {drow.get('grade')}"
            words = drow.get('words') or []
            if words:
                for i, word in enumerate(words):
                    ws_dist.append([
                        label if i == 0 else '',
                        drow.get('word_count', 0) if i == 0 else '',
                        f"{drow.get('word_pct', 0)}%" if i == 0 else '',
                        word,
                    ])
            else:
                ws_dist.append([label, drow.get('word_count', 0), f"{drow.get('word_pct', 0)}%", ''])
        ws_dist.column_dimensions['A'].width = 18
        ws_dist.column_dimensions['B'].width = 14
        ws_dist.column_dimensions['C'].width = 20
        ws_dist.column_dimensions['D'].width = 28

        # Sheet 3: New words (new to student) per grade
        ws2 = wb.create_sheet('New Words (new to student)')
        hdr_row(ws2, ['Standard (Std 1–N)','New word count','% new','Word'])
        for r in results:
            for i, word in enumerate(r.get('new_word_list', [])):
                ws2.append([
                    f"Std 1–{r['grade']}" if i==0 else '',
                    r['new_words'] if i==0 else '',
                    f"{r['new_pct']}%" if i==0 else '',
                    word
                ])
        ws2.column_dimensions['A'].width = 12
        ws2.column_dimensions['B'].width = 16
        ws2.column_dimensions['C'].width = 12
        ws2.column_dimensions['D'].width = 28

        # Sheet 4: New words per grade (same list, kept for detail)
        ws3 = wb.create_sheet('New Words Detail')
        hdr_row(ws3, ['Standard (Std 1–N)','New words (new to student)','Word'])
        for r in results:
            for i, word in enumerate(r.get('new_word_list', [])):
                ws3.append([
                    f"Std 1–{r['grade']}" if i==0 else '',
                    r['new_words'] if i==0 else '',
                    word
                ])
        ws3.column_dimensions['A'].width = 12
        ws3.column_dimensions['B'].width = 14
        ws3.column_dimensions['C'].width = 28

        # Sheet 5: Sentence analysis
        ws4 = wb.create_sheet('Sentence Analysis')
        ws4['A1'] = f'Target book: avg={tss.get("avg","—")} words/sent | max={tss.get("max","—")} | median={tss.get("median","—")} | sentences={tss.get("total_sentences",0):,}'
        ws4['A1'].font = Font(bold=True)
        ws4.append([])
        hdr_row(ws4, ['Standard','Grade Max Sentence','Grade Avg Sentence',
                      'Target Avg Sentence','Sentences Over Grade Max','% Over Max'], 3)
        for r in results:
            ws4.append([
                f"Std {r['grade']}",
                r.get('grade_sent_max','—'), r.get('grade_sent_avg','—'),
                tss.get('avg','—'),
                r.get('target_sentences_over_max','—'),
                f"{r.get('target_pct_over_max','—')}%",
            ])
        for col, w in enumerate([12,18,18,18,22,12], 1):
            ws4.column_dimensions[get_column_letter(col)].width = w

        # Sheet 5: Meaning-level appropriateness
        if meaning and meaning.get('enabled'):
            wsm = wb.create_sheet('Meaning Appropriateness')
            wsm['A1'] = f"Meaning-level target: Std {meaning.get('target_grade','—')} | Score: {meaning.get('appropriateness_pct','—')}% | Flags: {meaning.get('flagged_count',0)}"
            wsm['A1'].font = Font(bold=True)
            wsm.append([])
            hdr_row(wsm, ['Item','Type','Frequency','Detected level','Gap','Severity','Concept'], 3)
            for f in meaning.get('flagged', []):
                wsm.append([
                    f.get('item'), f.get('type'), f.get('freq'), f.get('level'),
                    f.get('gap'), f.get('severity'), f.get('concept')
                ])
            for col, w in enumerate([30,12,12,15,8,18,22], 1):
                wsm.column_dimensions[get_column_letter(col)].width = w

        # Sheet 5: Proper nouns
        ws5 = wb.create_sheet('Proper Nouns')
        ws5['A1'] = 'Proper nouns counted as known at all grade levels'
        ws5['A1'].font = Font(bold=True)
        ws5.append(['Word'])
        ws5['A2'].font = Font(bold=True)
        for w in proper_nouns: ws5.append([w])
        ws5.column_dimensions['A'].width = 30

        output = io.BytesIO()
        wb.save(output); output.seek(0)
        safe_name = re.sub(r'[^\w]','_', os.path.splitext(book_name)[0])
        return send_file(output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True, download_name=f'readability_{safe_name}.xlsx')
    except ImportError:
        return jsonify({'error': 'openpyxl not installed'}), 500

# ── Folder watcher integration ────────────────────────────────────────────────

import folder_watcher as _fw

def _get_loaded_hashes():
    """Return {filepath: file_hash} from DB for change detection."""
    conn = get_db()
    rows = conn.execute('SELECT filepath, file_hash FROM grade_files').fetchall()
    conn.close()
    return {r['filepath']: r['file_hash'] for r in rows}

def _watcher_process_fn(filepath, grade):
    return _process_grade_file(filepath, grade, source='folder')

@app.route('/api/config', methods=['GET'])
def get_config():
    cfg = _fw.load_config()
    return jsonify({**cfg, 'watcher_status': _fw.WATCHER_STATUS})

@app.route('/api/config', methods=['POST'])
def set_config():
    global WATCH_FOLDER
    data = request.json
    cfg = _fw.load_config()
    if 'watch_folder' in data:
        cfg['watch_folder'] = data['watch_folder']
    if 'mappings' in data:
        cfg['mappings'] = data['mappings']
    if 'auto_grade_from_name' in data:
        cfg['auto_grade_from_name'] = data['auto_grade_from_name']
    _fw.save_config(cfg)
    folder = cfg.get('watch_folder', '').strip()
    if folder and os.path.isdir(folder):
        WATCH_FOLDER = folder
        _fw.WATCHER_STATUS['folder'] = folder
        # Scan immediately in background
        t = threading.Thread(target=_fw.scan_folder,
                             args=(folder, cfg, _watcher_process_fn, _get_loaded_hashes),
                             daemon=True)
        t.start()
        _fw.start_watcher(folder, cfg, _watcher_process_fn, _get_loaded_hashes)
        return jsonify({'ok': True, 'message': f'Folder set to: {folder}. Scanning now…'})
    elif folder:
        return jsonify({'error': f'Folder not found: {folder}'}), 400
    return jsonify({'ok': True, 'message': 'Config saved'})

@app.route('/api/watcher/status')
def watcher_status():
    status = dict(_fw.WATCHER_STATUS)
    # Enrich current_files with elapsed time per file
    import time as _time
    now = _time.time()
    current = {}
    for fname, info in status.get('current_files', {}).items():
        current[fname] = {
            'stage':      info.get('stage', ''),
            'detail':     info.get('detail', ''),
            'elapsed_sec': round(now - info.get('started_at', now)),
        }
    status['current_files'] = current
    return jsonify(status)

@app.route('/api/watcher/scan', methods=['POST'])
def watcher_scan_now():
    """Manually trigger a rescan of the watch folder."""
    cfg = _fw.load_config()
    folder = cfg.get('watch_folder', '').strip()
    if not folder or not os.path.isdir(folder):
        return jsonify({'error': 'No valid watch folder configured'}), 400
    t = threading.Thread(target=_fw.scan_folder,
                         args=(folder, cfg, _watcher_process_fn, _get_loaded_hashes),
                         daemon=True)
    t.start()
    return jsonify({'ok': True, 'message': 'Scan started'})

# ── Word / Sentence Appropriateness Checker ──────────────────────────────────

@app.route('/api/check', methods=['POST'])
def check_words():
    """
    Paragraph / sentence checker.

    Runs the same comprehension analysis as the full book analyzer, but on
    typed or pasted text instead of a file.  Returns:
      - comprehension_table : per-grade known/unknown/% (same as book analysis)
      - best_grade          : first grade where comprehension >= 80 %
      - inline_tokens       : each token with its grade tag (for highlighting)
      - sentence_results    : per-sentence breakdown with comprehension %
      - unknown_list        : words not found in any grade
    """
    text = request.json.get('text', '').strip()
    if not text:
        return jsonify({'error': 'No text provided'}), 400

    raw_words = re.findall(r'[\u0B80-\u0BFF]{2,}', text)
    if not raw_words:
        return jsonify({'error': 'No Tamil words found in input'}), 400

    # ── Load grade data ───────────────────────────────────────────────────────
    conn = get_db()
    grade_rows  = conn.execute('SELECT grade, word FROM grade_words ORDER BY grade').fetchall()
    wgm_rows    = conn.execute('SELECT stem, first_grade FROM word_grade_map').fetchall()
    grade_metas = {r['grade']: dict(r) for r in
                   conn.execute('SELECT * FROM grade_meta ORDER BY grade').fetchall()}
    conn.close()

    if not grade_rows:
        return jsonify({'error': 'No school books loaded yet.'}), 400

    grade_vocab = {}
    for r in grade_rows:
        g = r['grade']
        if g not in grade_vocab: grade_vocab[g] = set()
        grade_vocab[g].add(r['word'])

    word_grade_map = {r['stem']: r['first_grade'] for r in wgm_rows}
    available_grades = sorted(grade_vocab.keys())

    # ── Stem every word, build unique set ─────────────────────────────────────
    stem_to_original = {}
    stem_freq        = {}
    all_stems        = []
    for w in raw_words:
        s = get_stem(w)
        all_stems.append(s)
        if s not in stem_to_original: stem_to_original[s] = w
        stem_freq[s] = stem_freq.get(s, 0) + 1

    unique_stems  = set(all_stems)
    total_unique  = len(unique_stems)

    # ── Comprehension table (same logic as book analysis) ─────────────────────
    comprehension_table = []
    cumulative = set()
    cumulative_prev = set()
    for g in range(1, 13):
        if g in grade_vocab: cumulative |= grade_vocab[g]
        if g not in available_grades: continue

        known   = unique_stems & cumulative
        unknown = unique_stems - known
        pct     = round(len(known) / total_unique * 100, 1) if total_unique else 0.0
        new_here = known - (unique_stems & cumulative_prev)

        gm      = grade_metas.get(g, {})
        gmax    = gm.get('sent_max', 0)

        comprehension_table.append({
            'grade':           g,
            'known':           len(known),
            'unknown':         len(unknown),
            'comprehension_pct': pct,
            'verdict':         ('Easy' if pct >= 90 else 'Readable' if pct >= 80
                                 else 'Challenging' if pct >= 60 else 'Very Hard'),
            'unknown_words':   sorted([stem_to_original.get(s, s) for s in unknown]),
            'new_words':       sorted([stem_to_original.get(s, s) for s in new_here]),
            'grade_sent_max':  gmax,
        })
        cumulative_prev = set(cumulative)

    best_grade = next(
        (r['grade'] for r in comprehension_table if r['comprehension_pct'] >= 80), None
    )

    # ── Inline token list for highlighted rendering ───────────────────────────
    # Tokenise the raw text preserving non-Tamil runs so the UI can
    # reconstruct the full paragraph with colour-coded Tamil words.
    inline_tokens = []
    pos = 0
    for m in re.finditer(r'[\u0B80-\u0BFF]{2,}', text):
        if m.start() > pos:
            inline_tokens.append({'type': 'text', 'value': text[pos:m.start()]})
        word  = m.group()
        stem  = get_stem(word)
        grade = word_grade_map.get(stem)
        inline_tokens.append({
            'type':  'word',
            'value': word,
            'stem':  stem,
            'grade': grade,
        })
        pos = m.end()
    if pos < len(text):
        inline_tokens.append({'type': 'text', 'value': text[pos:]})

    # ── Per-sentence breakdown ────────────────────────────────────────────────
    sentence_results = []
    sent_counts = sentence_word_counts(text)
    for sent in tokenize_sentences(text):
        sent_words = re.findall(r'[\u0B80-\u0BFF]{2,}', sent)
        if not sent_words: continue

        s_stems      = [get_stem(w) for w in sent_words]
        s_unique     = set(s_stems)
        s_total      = len(s_unique)
        word_count   = len(sent_words)

        # Per-grade comprehension for this sentence
        sent_grades = []
        s_cum = set()
        for g in range(1, 13):
            if g in grade_vocab: s_cum |= grade_vocab[g]
            if g not in available_grades: continue
            s_known   = s_unique & s_cum
            s_unknown = s_unique - s_known
            s_pct     = round(len(s_known) / s_total * 100, 1) if s_total else 0.0
            gmax      = grade_metas.get(g, {}).get('sent_max', 0)
            sent_grades.append({
                'grade':           g,
                'comprehension_pct': s_pct,
                'verdict':         ('Easy' if s_pct >= 90 else 'Readable' if s_pct >= 80
                                     else 'Challenging' if s_pct >= 60 else 'Very Hard'),
                'unknown_words':   sorted([stem_to_original.get(s, s) for s in s_unknown]),
                'length_ok':       word_count <= gmax if gmax else True,
            })

        best_sent_grade = next(
            (r['grade'] for r in sent_grades if r['comprehension_pct'] >= 80), None
        )
        unknown_in_sent = [
            stem_to_original.get(s, s) for s in s_unique
            if word_grade_map.get(s) is None
        ]

        sentence_results.append({
            'sentence':       sent,
            'word_count':     word_count,
            'best_grade':     best_sent_grade,
            'unknown_words':  unknown_in_sent,
            'grades':         sent_grades,
        })

    # ── Unknown words (not in any grade) ─────────────────────────────────────
    unknown_list = sorted(set(
        stem_to_original.get(s, s)
        for s in unique_stems
        if word_grade_map.get(s) is None
    ))

    return jsonify({
        'total_words':         len(raw_words),
        'unique_stems':        total_unique,
        'best_grade':          best_grade,
        'unknown_count':       len(unknown_list),
        'unknown_list':        unknown_list,
        'comprehension_table': comprehension_table,
        'inline_tokens':       inline_tokens,
        'sentence_results':    sentence_results,
    })


@app.route('/api/analytics/<int:analysis_id>')
def get_analytics(analysis_id):
    conn = get_db()
    row = conn.execute('SELECT analytics_json FROM analyses WHERE id = ?',
                       (analysis_id,)).fetchone()
    conn.close()
    if not row: return jsonify({'error': 'Not found'}), 404
    try:
        data = json.loads(row['analytics_json'] or '{}')
    except: data = {}
    return jsonify(data)


@app.route('/api/diagnose', methods=['POST'])
def diagnose_pdf():
    """
    Diagnostic endpoint: test what extract_text gets from a PDF.
    Accepts a filepath (must be on the server) and returns a full report.
    Only for debugging — not exposed in the main UI.
    """
    data = request.json or {}
    filepath = data.get('filepath', '').strip()

    if not filepath or not os.path.exists(filepath):
        return jsonify({'error': f'File not found: {filepath}'}), 400

    fname = os.path.basename(filepath)
    result = {
        'filename': fname,
        'filepath': filepath,
        'file_size_kb': round(os.path.getsize(filepath) / 1024, 1),
        'strategies': [],
    }

    import io as _io

    # Strategy 1: pdfminer page-by-page
    try:
        from pdfminer.pdfpage import PDFPage
        from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
        from pdfminer.converter import TextConverter
        from pdfminer.layout import LAParams

        rsrcmgr  = PDFResourceManager()
        pages_ok = 0; pages_err = 0; total_chars = 0; tamil_chars = 0
        sample   = ''

        with open(filepath, 'rb') as fh:
            for page_num, page in enumerate(PDFPage.get_pages(fh)):
                try:
                    buf = _io.StringIO()
                    dev = TextConverter(rsrcmgr, buf, laparams=LAParams())
                    PDFPageInterpreter(rsrcmgr, dev).process_page(page)
                    dev.close()
                    text = buf.getvalue()
                    total_chars += len(text)
                    tamil_chars += len(re.findall(r'[஀-௿]', text))
                    if not sample and text.strip():
                        sample = text.strip()[:120]
                    pages_ok += 1
                except Exception as e:
                    pages_err += 1

        result['strategies'].append({
            'name': 'pdfminer',
            'pages_ok': pages_ok,
            'pages_error': pages_err,
            'total_chars': total_chars,
            'tamil_chars': tamil_chars,
            'has_tamil': tamil_chars > 0,
            'sample': sample[:120] if sample else '',
            'sample_codepoints': [f'U+{ord(c):04X}' for c in (sample[:30] if sample else '') if ord(c) > 127][:15],
        })
    except Exception as e:
        result['strategies'].append({'name': 'pdfminer', 'error': str(e)})

    # Strategy 2: pdfplumber page-by-page
    try:
        import pdfplumber
        pages_ok2 = 0; pages_err2 = 0; total_chars2 = 0; tamil_chars2 = 0
        sample2 = ''

        with pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    t = page.extract_text() or ''
                    total_chars2 += len(t)
                    tamil_chars2 += len(re.findall(r'[஀-௿]', t))
                    if not sample2 and t.strip():
                        sample2 = t.strip()[:120]
                    pages_ok2 += 1
                except Exception as e:
                    pages_err2 += 1

        result['strategies'].append({
            'name': 'pdfplumber',
            'pages_ok': pages_ok2,
            'pages_error': pages_err2,
            'total_chars': total_chars2,
            'tamil_chars': tamil_chars2,
            'has_tamil': tamil_chars2 > 0,
            'sample': sample2[:120] if sample2 else '',
            'sample_codepoints': [f'U+{ord(c):04X}' for c in (sample2[:30] if sample2 else '') if ord(c) > 127][:15],
        })
    except Exception as e:
        result['strategies'].append({'name': 'pdfplumber', 'error': str(e)})

    # Verdict
    has_tamil = any(s.get('has_tamil') for s in result['strategies'])
    has_text   = any(s.get('total_chars', 0) > 0 for s in result['strategies'])
    if has_tamil:
        result['verdict'] = 'ok'
        result['verdict_msg'] = 'Tamil text found — should process correctly'
    elif has_text:
        result['verdict'] = 'encoding'
        result['verdict_msg'] = ('Text found but no Tamil Unicode. '
                                 'Font is likely TSCII/Bamini/TAB (pre-Unicode Tamil). '
                                 'Convert the PDF to Unicode Tamil using a tool like '
                                 'ilovepdf.com or Zamzar, or re-export from the source.')
    else:
        result['verdict'] = 'scanned'
        result['verdict_msg'] = ('No text layer found. PDF is a scanned image. '
                                 'Install tesseract with Tamil: '
                                 'sudo apt install tesseract-ocr tesseract-ocr-tam')

    return jsonify(result)


@app.route('/api/compare', methods=['POST'])
def compare_books():
    """
    Compare multiple analyses side-by-side.
    Input: { "ids": [1, 2, 3, ...] }  (up to 10)
    Returns per-book summary + per-grade comprehension matrix.
    """
    ids = (request.json or {}).get('ids', [])
    if not ids or len(ids) < 2:
        return jsonify({'error': 'Provide at least 2 analysis IDs'}), 400
    if len(ids) > 10:
        return jsonify({'error': 'Maximum 10 books at once'}), 400

    conn = get_db()
    books = []
    all_grades = set()

    for aid in ids:
        row = conn.execute('SELECT * FROM analyses WHERE id = ?', (aid,)).fetchone()
        if not row:
            continue
        row = dict(row)
        results = json.loads(row['results_json'] or '[]')
        analytics = json.loads(row.get('analytics_json') or 'null')

        # Build grade → comprehension lookup
        grade_comp = {r['grade']: r.get('known_pct', r.get('comprehension_pct', 0))
                      for r in results}
        grade_new  = {r['grade']: r.get('new_pct', 0) for r in results}
        all_grades.update(grade_comp.keys())

        # Best grade = first grade ≥ 80%
        best = next((r['grade'] for r in results
                     if r.get('known_pct', r.get('comprehension_pct', 0)) >= 80), None)

        # Sentence stats
        sent = json.loads(row.get('sentence_json') or '{}')
        tss  = sent.get('target', {})

        books.append({
            'id':           aid,
            'book_name':    row['book_name'],
            'analyzed_at':  row['analyzed_at'],
            'total_words':  row['total_words'],
            'unique_stems': row.get('unique_stems', 0),
            'best_grade':   best,
            'grade_comp':   grade_comp,
            'grade_new':    grade_new,
            'sent_avg':     tss.get('avg', 0),
            'sent_max':     tss.get('max', 0),
            # Analytics summary
            'ttr':          (analytics or {}).get('lexical', {}).get('ttr', 0),
            'dialogue_pct': (analytics or {}).get('dialogue', {}).get('dialogue_pct', 0),
            'readability_score': ((analytics or {}).get('readability_score') or {}).get('score'),
            'content_flags': len(((analytics or {}).get('content_flags') or {}).get('flags', [])),
        })

    conn.close()

    sorted_grades = sorted(all_grades)

    return jsonify({
        'books':  books,
        'grades': sorted_grades,
    })


@app.route('/api/batch/upload', methods=['POST'])
def batch_upload():
    """
    Accept multiple files for batch analysis.
    Returns a batch_id and list of pending extraction IDs.
    """
    files = request.files.getlist('files')
    if not files:
        return jsonify({'error': 'No files provided'}), 400

    conn = get_db()
    rows = conn.execute('SELECT word FROM grade_words').fetchall()
    conn.close()
    if not rows:
        return jsonify({'error': 'No school books uploaded yet.'}), 400

    batch_id  = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    results   = []
    os.makedirs('uploads', exist_ok=True)

    for file in files:
        if not file or not file.filename:
            continue
        fname    = werkzeug.utils.secure_filename(file.filename)
        filepath = os.path.join('uploads', f'batch_{batch_id}_{fname}')
        try:
            file.save(filepath)
        except Exception as e:
            results.append({'filename': fname, 'error': str(e)})
            continue

        # Run extraction synchronously (for now — could be made async)
        try:
            from app import _extract_stems_parallel, detect_proper_nouns, _db_write_lock
            stems, raw_unique, ss, text = _extract_stems_parallel(filepath)
            if stems is None:
                results.append({'filename': fname,
                                 'error': 'No Tamil text found. Scanned or non-Unicode PDF.'})
                os.unlink(filepath)
                continue

            stem_to_original = {}
            stem_freq        = {}
            for word in tokenize_tamil(text):
                s = get_stem(word)
                if s:
                    stem_freq[s]        = stem_freq.get(s, 0) + 1
                    stem_to_original[s] = stem_to_original.get(s, word)

            conn2 = get_db()
            gv_rows = conn2.execute('SELECT word FROM grade_words').fetchall()
            conn2.close()
            grade_vocab_union = {r['word'] for r in gv_rows}
            flagged = detect_proper_nouns(stem_freq, grade_vocab_union)

            tsc = sentence_word_counts(text)
            conn3 = get_db()
            # Note: do NOT delete pending_extractions here — batch mode needs all rows
            conn3.execute("""
                INSERT INTO pending_extractions
                  (book_name, created_at, total_words, unique_words, unique_stems,
                   stem_to_original, stem_freq_json, flagged_json, sentence_counts)
                VALUES (?,?,?,?,?,?,?,?,?)
            """, (fname, datetime.datetime.now().isoformat(),
                  sum(stem_freq.values()), raw_unique, len(stems),
                  json.dumps(stem_to_original), json.dumps(stem_freq),
                  json.dumps({s: list(r) for s, r in flagged.items()}),
                  json.dumps(tsc[:2000])))
            conn3.commit()
            pid = conn3.execute('SELECT last_insert_rowid()').fetchone()[0]
            conn3.close()

            results.append({
                'filename':    fname,
                'pending_id':  pid,
                'total_words': sum(stem_freq.values()),
                'unique_stems': len(stems),
                'proper_noun_count': len(flagged),
                'filepath':    filepath,
            })
        except Exception as e:
            results.append({'filename': fname, 'error': str(e)})
            if os.path.exists(filepath):
                os.unlink(filepath)

    return jsonify({'batch_id': batch_id, 'files': results})


@app.route('/api/batch/analyze', methods=['POST'])
def batch_analyze():
    """
    Run analysis on all pending IDs (from batch_upload) skipping proper-noun review.
    Input: { "pending_ids": [1, 2, ...] }
    Returns list of { pending_id, analysis_id, book_name, best_grade, error }.
    """
    pending_ids = (request.json or {}).get('pending_ids', [])
    if not pending_ids:
        return jsonify({'error': 'No pending IDs provided'}), 400

    results = []
    for pid in pending_ids:
        try:
            d = _run_single_analysis(pid, confirmed_proper_stems=[])
            results.append({
                'pending_id':  pid,
                'analysis_id': d.get('analysis_id'),
                'book_name':   d.get('book_name'),
                'best_grade':  d.get('best_grade'),
                'error':       d.get('error'),
            })
        except Exception as e:
            results.append({'pending_id': pid, 'error': str(e)})

    return jsonify({'results': results})


@app.route('/api/card/<int:analysis_id>')
def generate_card(analysis_id):
    """
    Generate a one-page A5 reading-level summary card as PDF.
    Clean, printable, shareable — designed for publishers and teachers.
    """
    conn = get_db()
    row  = conn.execute('SELECT * FROM analyses WHERE id = ?', (analysis_id,)).fetchone()
    conn.close()
    if not row: return jsonify({'error': 'Not found'}), 404

    row      = dict(row)
    results  = json.loads(row['results_json'] or '[]')
    book_name = row['book_name']
    analytics = json.loads(row.get('analytics_json') or 'null') or {}

    best_grade = next((r['grade'] for r in results
                       if r.get('known_pct', r.get('comprehension_pct', 0)) >= 80), None)
    best_pct   = next((r.get('known_pct', r.get('comprehension_pct', 0))
                       for r in results if r['grade'] == best_grade), 0) if best_grade else 0

    try:
        from reportlab.lib.pagesizes import A5
        from reportlab.lib import colors
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import cm, mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # Font setup (same as main report)
        _base = os.path.dirname(os.path.abspath(__file__))
        TAMIL_FONT = 'Helvetica'; TAMIL_BOLD = 'Helvetica-Bold'
        for fp in [os.path.join(_base,'fonts','FreeSerif.ttf'),
                   '/usr/share/fonts/truetype/freefont/FreeSerif.ttf']:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont('CardTamil', fp))
                    TAMIL_FONT = 'CardTamil'
                    break
                except: pass
        for fp in [os.path.join(_base,'fonts','FreeSerifBold.ttf'),
                   '/usr/share/fonts/truetype/freefont/FreeSerifBold.ttf']:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont('CardTamilBold', fp))
                    TAMIL_BOLD = 'CardTamilBold'
                    break
                except: pass

        buf = io.BytesIO()
        W, H = A5
        doc = SimpleDocTemplate(buf, pagesize=A5,
                                 leftMargin=1.2*cm, rightMargin=1.2*cm,
                                 topMargin=1.2*cm, bottomMargin=1.2*cm)

        # ── Colour palette ──────────────────────────────────────────────────
        TEAL    = colors.HexColor('#1D9E75')
        TEAL_LT = colors.HexColor('#E1F5EE')
        AMBER   = colors.HexColor('#EF9F27')
        RED     = colors.HexColor('#E24B4A')
        GREY    = colors.HexColor('#5c574f')
        LGREY   = colors.HexColor('#f5f5f3')
        WHITE   = colors.white
        BLACK   = colors.HexColor('#1a1a18')

        grade_color = TEAL if best_grade and best_grade <= 5 else AMBER if best_grade else RED

        def S(name, **kw):
            return ParagraphStyle(name, **kw)

        sTitle  = S('ct', fontName=TAMIL_BOLD,  fontSize=13, leading=17,
                    alignment=TA_CENTER, textColor=BLACK, spaceAfter=2)
        sSub    = S('cs', fontName=TAMIL_FONT,  fontSize=8.5, leading=12,
                    alignment=TA_CENTER, textColor=GREY, spaceAfter=8)
        sLabel  = S('cl', fontName='Helvetica-Bold', fontSize=7,
                    textColor=GREY, spaceBefore=6, spaceAfter=1)
        sNote   = S('cn', fontName=TAMIL_FONT,  fontSize=7.5, leading=11,
                    textColor=GREY, spaceAfter=4)
        sFlag   = S('cf', fontName=TAMIL_FONT,  fontSize=8, leading=12,
                    textColor=BLACK)

        story = []

        # ── Header ──────────────────────────────────────────────────────────
        story.append(Table([[Paragraph('<b>READING LEVEL CARD</b>',
                              S('bh', fontName='Helvetica-Bold', fontSize=10,
                                alignment=TA_CENTER, textColor=WHITE))]],
                           colWidths=[W - 2.4*cm],
                           style=TableStyle([
                               ('BACKGROUND', (0,0),(-1,-1), TEAL),
                               ('TOPPADDING',  (0,0),(-1,-1), 8),
                               ('BOTTOMPADDING',(0,0),(-1,-1), 8),
                           ])))
        story.append(Spacer(1, 6))

        # Book title
        display_name = os.path.splitext(book_name)[0].replace('_',' ')
        story.append(Paragraph(display_name, sTitle))
        story.append(Paragraph(
            f"Analyzed on {row['analyzed_at'][:10]}  ·  "
            f"{row['total_words']:,} words  ·  {row.get('unique_stems',0):,} unique stems",
            sSub))
        story.append(HRFlowable(width='100%', thickness=0.5,
                                 color=colors.HexColor('#cccccc'), spaceAfter=8))

        # ── Best-fit grade (big badge) ───────────────────────────────────────
        grade_label = f'Standard {best_grade}' if best_grade else 'Beyond Std 12'
        pct_label   = f'{best_pct:.0f}% comprehension' if best_grade else 'Needs advanced vocabulary'
        story.append(Table([[
            Paragraph(f'<b>{grade_label}</b>',
                      S('bg', fontName=TAMIL_BOLD, fontSize=22,
                        alignment=TA_CENTER, textColor=WHITE)),
            Paragraph(pct_label,
                      S('bp', fontName=TAMIL_FONT, fontSize=9,
                        alignment=TA_CENTER, textColor=WHITE)),
        ]], colWidths=[(W-2.4*cm)*0.55, (W-2.4*cm)*0.45],
            style=TableStyle([
                ('BACKGROUND',   (0,0),(-1,-1), grade_color),
                ('TOPPADDING',   (0,0),(-1,-1), 10),
                ('BOTTOMPADDING',(0,0),(-1,-1), 10),
                ('LEFTPADDING',  (0,0),(-1,-1), 8),
                ('RIGHTPADDING', (0,0),(-1,-1), 8),
                ('VALIGN',       (0,0),(-1,-1), 'MIDDLE'),
            ])))
        story.append(Spacer(1, 8))

        # ── Comprehension bar chart ──────────────────────────────────────────
        story.append(Paragraph('COMPREHENSION BY STANDARD', sLabel))
        CW = W - 2.4*cm

        bar_rows = [[
            Paragraph(f'<b>Std {r["grade"]}</b>',
                      S(f'g{r["grade"]}', fontName='Helvetica-Bold', fontSize=7,
                        textColor=BLACK)),
            Table([['']], colWidths=[
                (CW*0.55) * min(r.get("known_pct", r.get("comprehension_pct",0)), 100) / 100
            ], rowHeights=[10],
            style=TableStyle([
                ('BACKGROUND',(0,0),(-1,-1),
                 TEAL if r.get("known_pct",r.get("comprehension_pct",0))>=80
                 else AMBER if r.get("known_pct",r.get("comprehension_pct",0))>=60
                 else RED),
                ('TOPPADDING',(0,0),(-1,-1),0),
                ('BOTTOMPADDING',(0,0),(-1,-1),0),
                ('LEFTPADDING',(0,0),(-1,-1),0),
            ])) if r.get("known_pct",r.get("comprehension_pct",0)) > 0 else Paragraph('',sNote),
            Paragraph(f'{r.get("known_pct",r.get("comprehension_pct",0)):.0f}%',
                      S(f'p{r["grade"]}', fontName='Helvetica', fontSize=7, textColor=GREY)),
        ] for r in results]

        if bar_rows:
            story.append(Table(bar_rows,
                               colWidths=[1.1*cm, CW*0.56, 0.9*cm],
                               style=TableStyle([
                                   ('TOPPADDING',   (0,0),(-1,-1), 3),
                                   ('BOTTOMPADDING',(0,0),(-1,-1), 3),
                                   ('LEFTPADDING',  (0,0),(-1,-1), 0),
                                   ('VALIGN',       (0,0),(-1,-1), 'MIDDLE'),
                               ])))
        story.append(Spacer(1, 8))

        # ── Quick stats row ─────────────────────────────────────────────────
        lex  = analytics.get('lexical', {})
        wl   = analytics.get('word_length', {})
        dial = analytics.get('dialogue', {})
        cf   = analytics.get('content_flags', {})
        rs   = analytics.get('readability_score', {}) or {}

        sent_data = json.loads(row.get('sentence_json') or '{}')
        tss = sent_data.get('target', {})

        stats = [
            ['Avg sentence', f"{tss.get('avg',0):.1f} words"],
            ['Vocab diversity', f"TTR {lex.get('ttr',0):.0f}%"],
            ['Dialogue', f"{dial.get('dialogue_pct',0):.0f}%"],
            ['Difficulty score', f"{rs.get('score','—')}"],
        ]

        story.append(Paragraph('QUICK STATS', sLabel))
        story.append(Table(
            [[Paragraph(f'<b>{k}</b>', S('sk',fontName='Helvetica-Bold',fontSize=7,textColor=GREY)),
              Paragraph(v, S('sv',fontName='Helvetica',fontSize=8,textColor=BLACK))]
             for k, v in stats],
            colWidths=[CW*0.45, CW*0.55],
            style=TableStyle([
                ('BACKGROUND',   (0,0),(-1,-1), LGREY),
                ('TOPPADDING',   (0,0),(-1,-1), 4),
                ('BOTTOMPADDING',(0,0),(-1,-1), 4),
                ('LEFTPADDING',  (0,0),(-1,-1), 6),
                ('ROWBACKGROUNDS',(0,0),(-1,-1),[LGREY, WHITE]),
            ])))
        story.append(Spacer(1, 8))

        # ── Content flags ───────────────────────────────────────────────────
        flags = cf.get('flags', [])
        if flags:
            story.append(Paragraph('CONTENT FLAGS', sLabel))
            flag_items = []
            for fl in flags[:4]:
                sev_color = RED if fl['severity']=='warning' else AMBER
                flag_items.append([
                    Paragraph(f"● {fl['category']}",
                              S('fi',fontName=TAMIL_FONT,fontSize=7.5,textColor=sev_color)),
                    Paragraph(f"Age {fl['min_age']}+",
                              S('fa',fontName='Helvetica',fontSize=7,textColor=GREY)),
                ])
            if flag_items:
                story.append(Table(flag_items,
                                   colWidths=[CW*0.75, CW*0.25],
                                   style=TableStyle([
                                       ('TOPPADDING',   (0,0),(-1,-1), 3),
                                       ('BOTTOMPADDING',(0,0),(-1,-1), 3),
                                       ('LEFTPADDING',  (0,0),(-1,-1), 0),
                                   ])))
            story.append(Spacer(1, 4))

        # ── Footer ──────────────────────────────────────────────────────────
        story.append(HRFlowable(width='100%', thickness=0.5,
                                 color=colors.HexColor('#cccccc'), spaceBefore=4, spaceAfter=4))
        story.append(Paragraph(
            'Generated by Tamil Book Readability Analyzer (v28) · All analysis local & offline',
            S('ft', fontName='Helvetica', fontSize=6.5,
              alignment=TA_CENTER, textColor=colors.HexColor('#aaaaaa'))))

        doc.build(story)
        buf.seek(0)
        safe_name = os.path.splitext(book_name)[0][:40]
        return buf.getvalue(), 200, {
            'Content-Type':        'application/pdf',
            'Content-Disposition': f'attachment; filename="card_{safe_name}.pdf"',
        }

    except Exception as e:
        import traceback
        logging.getLogger('app').error(f'Card generation failed: {e}')
        logging.getLogger('app').debug(traceback.format_exc())
        return jsonify({'error': f'Card generation failed: {e}'}), 500


# ── Load simplifier module ────────────────────────────────────────────────────
import simplifier as _simplifier

def _build_simplifier_engine():
    """Build a SimplifierEngine from the current grade DB."""
    conn = get_db()
    grade_rows = conn.execute(
        'SELECT grade, word FROM grade_words ORDER BY grade'
    ).fetchall()
    wgm_rows = conn.execute(
        'SELECT stem, first_grade FROM word_grade_map'
    ).fetchall()
    # Build grade frequency from grade_words (each occurrence = 1 for now)
    freq_rows = conn.execute(
        'SELECT word, COUNT(*) as cnt FROM grade_words GROUP BY word'
    ).fetchall()
    conn.close()

    grade_vocab = {}
    for r in grade_rows:
        g = r['grade']
        if g not in grade_vocab:
            grade_vocab[g] = set()
        grade_vocab[g].add(r['word'])

    word_grade_map = {r['stem']: r['first_grade'] for r in wgm_rows}
    grade_freq     = {r['word']: r['cnt'] for r in freq_rows}

    return _simplifier.SimplifierEngine(
        grade_vocab    = grade_vocab,
        word_grade_map = word_grade_map,
        stem_fn        = get_stem,
        grade_freq     = grade_freq,
    )


@app.route('/api/simplify', methods=['POST'])
def simplify_text():
    """
    Simplify a Tamil text for a target grade.
    Input: multipart/form-data with:
      - file (PDF/DOCX/TXT)  OR  text (plain Tamil string)
      - target_grade (int)
    Returns structured simplification report.
    """
    target_grade = int(request.form.get('target_grade') or
                       (request.json or {}).get('target_grade') or 3)

    # Get text from file upload or direct text input
    text = ''
    if 'file' in request.files:
        file = request.files['file']
        if file and file.filename:
            fname    = werkzeug.utils.secure_filename(file.filename)
            filepath = os.path.join('uploads', f'simplify_{fname}')
            os.makedirs('uploads', exist_ok=True)
            file.save(filepath)
            try:
                text = extract_text(filepath)
            finally:
                if os.path.exists(filepath):
                    os.unlink(filepath)
    elif request.is_json:
        text = (request.json or {}).get('text', '')
    else:
        text = request.form.get('text', '')

    if not text or not text.strip():
        return jsonify({'error': 'No Tamil text provided or extracted.'}), 400

    # Check grade DB is loaded
    conn = get_db()
    wc = conn.execute('SELECT COUNT(*) FROM word_grade_map').fetchone()[0]
    conn.close()
    if not wc:
        return jsonify({'error': 'No school books loaded yet. Load grade books first.'}), 400

    try:
        engine = _build_simplifier_engine()
        report = engine.simplify_text(text, target_grade)
        return jsonify(report)
    except Exception as e:
        logging.getLogger('app').error(f'Simplify failed: {e}')
        return jsonify({'error': str(e)}), 500


@app.route('/api/simplify/export', methods=['POST'])
def simplify_export():
    """
    Export the simplification report as a Word document.
    Original text with hard words highlighted + suggested rewrites.
    Input JSON: the report dict returned by /api/simplify
    """
    report = request.json
    if not report:
        return jsonify({'error': 'No report data provided'}), 400

    target_grade = report.get('target_grade', 3)
    sentences    = report.get('sentences', [])

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn
        from docx.enum.text import WD_COLOR_INDEX
        import io as _io

        doc = Document()
        # Page setup
        section = doc.sections[0]
        section.page_width  = Inches(8.27)
        section.page_height = Inches(11.69)

        # Title
        title = doc.add_heading(level=1)
        title.clear()
        run = title.add_run(f'Tamil Text Simplification Report — Target: Standard {target_grade}')
        run.font.size = Pt(14)

        # Summary
        summary = doc.add_paragraph()
        summary.add_run(
            f'Total sentences: {report.get("total_sentences",0)}  |  '
            f'Total Tamil words: {report.get("total_words",0)}  |  '
            f'Hard words: {report.get("hard_word_count",0)} '
            f'({report.get("hard_word_pct",0):.1f}%)  |  '
            f'Complex sentences: {report.get("complex_sentence_count",0)}'
        ).font.size = Pt(10)
        doc.add_paragraph()

        final_text = (report.get('final_text') or '').strip()
        if final_text:
            doc.add_heading('Edited simplified draft', level=2)
            for para in final_text.splitlines():
                if para.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(para.strip())
                    r.font.size = Pt(11)
            doc.add_page_break()

        RED    = RGBColor(0xC0, 0x39, 0x2B)
        AMBER  = RGBColor(0xB8, 0x7A, 0x12)
        GREEN  = RGBColor(0x0F, 0x6E, 0x56)
        GREY   = RGBColor(0x88, 0x88, 0x88)

        for i, sent in enumerate(sentences, 1):
            if not sent.get('is_complex'):
                # Simple sentence — just write it in grey
                p = doc.add_paragraph()
                r = p.add_run(sent['original'])
                r.font.size  = Pt(11)
                r.font.color.rgb = GREY
                continue

            # ── Original sentence with hard words highlighted ──────────────
            doc.add_heading(f'Sentence {i}', level=3)

            orig_para = doc.add_paragraph()
            orig_para.paragraph_format.space_after = Pt(2)
            original = sent['original']
            hard_set = {h['word'] for h in sent.get('hard_words', [])}

            # Tokenise and write word by word with highlighting
            pos = 0
            for m in re.finditer(r'[஀-௿]{2,}', original):
                # Non-Tamil before this word
                if m.start() > pos:
                    r = orig_para.add_run(original[pos:m.start()])
                    r.font.size = Pt(11)
                w = m.group()
                r = orig_para.add_run(w)
                r.font.size = Pt(11)
                if w in hard_set:
                    r.font.color.rgb = RED
                    r.font.bold = True
                pos = m.end()
            if pos < len(original):
                r = orig_para.add_run(original[pos:])
                r.font.size = Pt(11)

            # ── Word replacements ──────────────────────────────────────────
            if sent.get('hard_words'):
                doc.add_paragraph('Word suggestions:', style='List Bullet')
                for h in sent['hard_words']:
                    sug_para = doc.add_paragraph(style='List Bullet 2')
                    r = sug_para.add_run(h['word'])
                    r.font.color.rgb = RED
                    r.font.bold = True
                    r.font.size = Pt(10)
                    grade_label = f" (Std {h['grade']})" if h['grade'] else " (unknown grade)"
                    sug_para.add_run(grade_label).font.size = Pt(9)
                    sug_para.add_run('  →  ').font.size = Pt(10)
                    if h['suggestions']:
                        # Best suggestion in green
                        best = h['suggestions'][0]
                        rb = sug_para.add_run(best['stem'])
                        rb.font.color.rgb = GREEN
                        rb.font.bold = True
                        rb.font.size = Pt(10)
                        rb2 = sug_para.add_run(f" (Std {best['grade']})")
                        rb2.font.size = Pt(9)
                        # Other options in amber
                        for alt in h['suggestions'][1:]:
                            sug_para.add_run(',  ').font.size = Pt(10)
                            ra = sug_para.add_run(alt['stem'])
                            ra.font.color.rgb = AMBER
                            ra.font.size = Pt(10)
                            ra2 = sug_para.add_run(f" (Std {alt['grade']})")
                            ra2.font.size = Pt(9)
                    else:
                        r_na = sug_para.add_run('No suggestion found — manual rewrite needed')
                        r_na.font.color.rgb = AMBER
                        r_na.font.size = Pt(9)

            # ── Rewritten sentence ─────────────────────────────────────────
            if sent.get('rewritten') and sent['rewritten'] != sent['original']:
                rw_para = doc.add_paragraph()
                rw_para.add_run('Suggested rewrite:  ').font.size = Pt(9)
                rr = rw_para.add_run(sent['rewritten'])
                rr.font.size  = Pt(11)
                rr.font.color.rgb = GREEN

            # ── Split suggestion ───────────────────────────────────────────
            if sent.get('split_suggestion'):
                sp_para = doc.add_paragraph()
                sp_para.add_run('Split into shorter sentences:  ').font.size = Pt(9)
                for part in sent['split_suggestion']:
                    rp = sp_para.add_run(part + '  ')
                    rp.font.size  = Pt(11)
                    rp.font.color.rgb = RGBColor(0x1A, 0x5F, 0xA8)

            doc.add_paragraph()  # spacer between sentences

        buf = _io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue(), 200, {
            'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'Content-Disposition': f'attachment; filename="simplification_std{target_grade}.docx"',
        }

    except Exception as e:
        logging.getLogger('app').error(f'Simplify export failed: {e}')
        return jsonify({'error': str(e)}), 500


# ── Word Library ──────────────────────────────────────────────────────────────
import word_library as _wlib

_WLIB_BUILD_STATUS = {'running': False, 'progress': 0, 'message': '', 'last': None}
_WLIB_STATUS_LOCK  = threading.Lock()

def _wlib_progress(done, total, stage):
    with _WLIB_STATUS_LOCK:
        pct = round(done / max(total, 1) * 100, 1)
        _WLIB_BUILD_STATUS.update({'progress': pct, 'message': stage})


@app.route('/api/library/stats')
def library_stats():
    try:
        stats = _wlib.get_stats()
        return jsonify(stats)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/library/search')
def library_search():
    q     = request.args.get('q', '').strip()
    limit = int(request.args.get('limit', 20))
    if not q:
        return jsonify([])
    try:
        return jsonify(_wlib.search_word(q, limit=limit))
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/library/grade/<int:grade>')
def library_by_grade(grade):
    concept  = request.args.get('concept', 'all')
    limit    = int(request.args.get('limit', 200))
    offset   = int(request.args.get('offset', 0))
    confirmed = request.args.get('confirmed') == '1'
    try:
        return jsonify(_wlib.get_by_grade(
            grade, concept=concept,
            confirmed_only=confirmed,
            limit=limit, offset=offset,
        ))
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/library/build', methods=['POST'])
def library_build():
    """
    Build/refresh the word library from textbooks already loaded in the main DB.
    Fast — runs synchronously.
    """
    with _WLIB_STATUS_LOCK:
        if _WLIB_BUILD_STATUS['running']:
            return jsonify({'error': 'Build already running'}), 409

    try:
        result = _wlib.build_from_textbooks(
            main_db_path=DB_PATH,
            stem_fn=get_stem,
            progress_cb=_wlib_progress,
        )
        with _WLIB_STATUS_LOCK:
            _WLIB_BUILD_STATUS['last'] = result
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/library/import_book', methods=['POST'])
def library_import_book():
    """
    Import a children's book PDF/DOCX/TXT into the word library.
    Accepts: file + grade (optional) + source_name (optional).
    """
    file        = request.files.get('file')
    grade_hint  = request.form.get('grade')
    source_name = request.form.get('source_name', '').strip()

    if not file or not file.filename:
        return jsonify({'error': 'No file provided'}), 400

    grade_int = int(grade_hint) if grade_hint and grade_hint.isdigit() else None
    fname     = werkzeug.utils.secure_filename(file.filename)
    if not source_name:
        source_name = os.path.splitext(fname)[0][:60]

    filepath = os.path.join('uploads', f'lib_{fname}')
    os.makedirs('uploads', exist_ok=True)
    file.save(filepath)

    try:
        # Load known grade map for inference
        conn = get_db()
        wgm  = {r['stem']: r['first_grade']
                for r in conn.execute('SELECT stem, first_grade FROM word_grade_map').fetchall()}
        conn.close()

        result = _wlib.import_from_book(
            filepath=filepath,
            source_name=source_name,
            grade_hint=grade_int,
            extract_fn=extract_text,
            stem_fn=get_stem,
            known_grade_map=wgm,
            progress_cb=_wlib_progress,
        )
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        if os.path.exists(filepath):
            os.unlink(filepath)


@app.route('/api/library/wiki_download', methods=['POST'])
def library_wiki_download():
    """
    Start background download of Tamil Wikipedia/Wiktionary dump.
    dump_type: abstracts | full | wiktionary
    """
    dump_type = (request.json or {}).get('dump_type', 'abstracts')

    def _bg():
        with _WLIB_STATUS_LOCK:
            _WLIB_BUILD_STATUS['running'] = True
            _WLIB_BUILD_STATUS['message'] = f'Downloading {dump_type}...'

        def _prog(pct, msg):
            with _WLIB_STATUS_LOCK:
                _WLIB_BUILD_STATUS.update({'progress': pct, 'message': msg})

        path = _wlib.download_wiki_dump(dump_type, progress_cb=_prog)

        if path:
            # Auto-import after download
            with _WLIB_STATUS_LOCK:
                _WLIB_BUILD_STATUS['message'] = 'Processing Wikipedia text...'
            conn = get_db()
            wgm  = {r['stem']: r['first_grade']
                    for r in conn.execute('SELECT stem, first_grade FROM word_grade_map').fetchall()}
            conn.close()
            result = _wlib.import_from_wiki_dump(
                path, stem_fn=get_stem,
                known_grade_map=wgm,
                progress_cb=_wlib_progress,
            )
            with _WLIB_STATUS_LOCK:
                _WLIB_BUILD_STATUS.update({
                    'running': False, 'progress': 100,
                    'message': f'Done: {result.get("added",0)} added, {result.get("updated",0)} updated',
                    'last': result,
                })
        else:
            with _WLIB_STATUS_LOCK:
                _WLIB_BUILD_STATUS.update({
                    'running': False, 'message': 'Download failed — check internet connection',
                })

    t = threading.Thread(target=_bg, daemon=True)
    t.start()
    return jsonify({'ok': True, 'message': f'Downloading {dump_type} in background...'})


@app.route('/api/library/wiki_status')
def library_wiki_status():
    with _WLIB_STATUS_LOCK:
        return jsonify(dict(_WLIB_BUILD_STATUS))


@app.route('/api/library/manual', methods=['POST'])
def library_manual_entry():
    """Add or override a word manually."""
    data = request.json or {}
    stem  = (data.get('stem') or '').strip()
    word  = (data.get('display_word') or stem).strip()
    grade = data.get('grade_level')

    if not stem or not grade:
        return jsonify({'error': 'stem and grade_level are required'}), 400

    try:
        result = _wlib.manual_entry(
            stem=stem, display_word=word,
            grade_level=int(grade),
            definition=data.get('definition'),
            part_of_speech=data.get('part_of_speech'),
            example=data.get('example'),
        )
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/library/export')
def library_export():
    """Export the word library as an Excel file."""
    try:
        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tf:
            path = tf.name
        _wlib.export_to_excel(path)
        with open(path, 'rb') as f:
            data = f.read()
        os.unlink(path)
        return data, 200, {
            'Content-Type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'Content-Disposition': 'attachment; filename="tamil_word_library.xlsx"',
        }
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/library/sync_to_analyzer', methods=['POST'])
def library_sync_to_analyzer():
    """
    Sync confirmed library words back into the main analyzer DB.
    Only teacher-confirmed words (confirmed=1) with a valid grade are synced.
    This enriches the readability analysis with the broader vocabulary.
    """
    grade_filter = (request.json or {}).get('grade')

    lib_conn = _wlib.get_lib_db()
    if grade_filter:
        rows = lib_conn.execute(
            'SELECT stem, grade_level FROM word_library WHERE confirmed=1 AND grade_level=?',
            (int(grade_filter),)
        ).fetchall()
    else:
        rows = lib_conn.execute(
            'SELECT stem, grade_level FROM word_library WHERE confirmed=1'
        ).fetchall()
    lib_conn.close()

    if not rows:
        return jsonify({'message': 'No confirmed words to sync', 'synced': 0})

    conn = get_db()
    synced = 0
    CHUNK  = 2000
    rows_list = [(r['grade_level'], r['stem']) for r in rows]

    for i in range(0, len(rows_list), CHUNK):
        chunk = rows_list[i:i+CHUNK]
        conn.executemany(
            'INSERT OR IGNORE INTO grade_words (grade, word) VALUES (?, ?)',
            chunk
        )
        conn.executemany('''
            INSERT INTO word_grade_map (stem, first_grade) VALUES (?, ?)
            ON CONFLICT(stem) DO UPDATE
              SET first_grade = MIN(first_grade, excluded.first_grade)
        ''', [(stem, grade) for grade, stem in chunk])
        synced += len(chunk)

    conn.commit()
    conn.close()
    return jsonify({'synced': synced, 'message': f'{synced} words synced to analyzer'})



@app.route('/api/v27/home_metrics')
def v27_home_metrics():
    try:
        return jsonify(_v27.collect_home_metrics(DB_PATH, 'data/meaning_kb', 'data/cache'))
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/v27/analysis/<int:analysis_id>')
def v27_analysis_details(analysis_id):
    conn = get_db()
    row = conn.execute('SELECT * FROM analyses WHERE id=?', (analysis_id,)).fetchone()
    conn.close()
    if not row:
        return jsonify({'error': 'Analysis not found'}), 404
    try:
        if 'v27_json' in row.keys() and row['v27_json']:
            return jsonify(json.loads(row['v27_json']))
    except Exception:
        pass
    try:
        results = json.loads(row['results_json'] or '[]')
        sentence = json.loads(row['sentence_json'] or '{}')
        meaning = json.loads(row['meaning_json'] or '{}') if 'meaning_json' in row.keys() and row['meaning_json'] else {}
        suitability = json.loads(row['suitability_json'] or '{}') if 'suitability_json' in row.keys() and row['suitability_json'] else {}
        return jsonify(_v27.build_offline_intelligence('', results, sentence.get('target_counts', []), meaning, suitability, 'data/meaning_kb'))
    except Exception as e:
        return jsonify({'enabled': False, 'error': str(e)}), 500




# ─────────────────────────────────────────────────────────────
# v28 Optional Local LLM / Ollama endpoints
# Default is local-only and safe to disable. No paid API required.
# ─────────────────────────────────────────────────────────────
def _ai_config():
    try:
        cfg = _fw.load_config() if '_fw' in globals() else {}
    except Exception:
        cfg = {}
    ai = cfg.get('ai', {}) if isinstance(cfg, dict) else {}
    return {
        'enabled': bool(ai.get('enabled', False)),
        'provider': ai.get('provider', 'ollama'),
        'base_url': ai.get('base_url', DEFAULT_BASE_URL),
        'model': ai.get('model', DEFAULT_MODEL),
    }


def _save_ai_config(new_ai):
    try:
        cfg = _fw.load_config() if '_fw' in globals() else {}
    except Exception:
        cfg = {}
    if not isinstance(cfg, dict):
        cfg = {}
    cfg['ai'] = new_ai
    try:
        _fw.save_config(cfg)
    except Exception:
        # fallback to config.json in project root
        import json as _json
        with open('config.json', 'w', encoding='utf-8') as f:
            _json.dump(cfg, f, ensure_ascii=False, indent=2)


@app.route('/api/ai/status')
def api_ai_status():
    cfg = _ai_config()
    health = ollama_health(cfg['base_url']) if cfg['provider'] == 'ollama' else {'available': False, 'models': [], 'error': 'Unsupported provider'}
    return jsonify({
        'enabled': cfg['enabled'],
        'provider': cfg['provider'],
        'base_url': cfg['base_url'],
        'model': cfg['model'],
        'available': health.get('available', False),
        'models': health.get('models', []),
        'error': health.get('error'),
        'recommended': ['qwen2.5:7b-instruct', 'qwen2.5:3b-instruct', 'llama3.1:8b']
    })


@app.route('/api/ai/settings', methods=['POST'])
def api_ai_settings():
    data = request.get_json(silent=True) or {}
    enabled = bool(data.get('enabled', False))
    model = (data.get('model') or DEFAULT_MODEL).strip()
    base_url = (data.get('base_url') or DEFAULT_BASE_URL).strip().rstrip('/')
    provider = (data.get('provider') or 'ollama').strip()
    cfg = {'enabled': enabled, 'provider': provider, 'model': model, 'base_url': base_url}
    _save_ai_config(cfg)
    return jsonify({'ok': True, **cfg})


def _require_ai():
    cfg = _ai_config()
    if not cfg['enabled']:
        return None, (jsonify({'error': 'AI assistant is disabled in Settings. Enable Local Ollama first.'}), 400)
    health = ollama_health(cfg['base_url'])
    if not health.get('available'):
        return None, (jsonify({'error': 'Ollama is not reachable. Start Ollama and pull the selected model.', 'details': health.get('error')}), 503)
    return cfg, None


@app.route('/api/ai/rewrite', methods=['POST'])
def api_ai_rewrite():
    cfg, err = _require_ai()
    if err: return err
    data = request.get_json(silent=True) or {}
    text = (data.get('text') or '').strip()
    grade = int(data.get('grade') or data.get('target_grade') or 3)
    if not text:
        return jsonify({'error': 'No text provided'}), 400
    out = tamil_author_rewrite(text[:6000], target_grade=grade, model=cfg['model'], base_url=cfg['base_url'])
    return jsonify({'ok': True, 'mode': 'local_ollama', 'model': cfg['model'], 'result': out})


@app.route('/api/ai/explain', methods=['POST'])
def api_ai_explain():
    cfg, err = _require_ai()
    if err: return err
    data = request.get_json(silent=True) or {}
    text = (data.get('text') or '').strip()
    grade = int(data.get('grade') or data.get('target_grade') or 3)
    if not text:
        return jsonify({'error': 'No text provided'}), 400
    out = tamil_simple_explanation(text[:4000], target_grade=grade, model=cfg['model'], base_url=cfg['base_url'])
    return jsonify({'ok': True, 'mode': 'local_ollama', 'model': cfg['model'], 'result': out})


@app.route('/api/ai/lesson_plan', methods=['POST'])
def api_ai_lesson_plan():
    cfg, err = _require_ai()
    if err: return err
    data = request.get_json(silent=True) or {}
    words = data.get('words') or []
    concepts = data.get('concepts') or []
    grade = int(data.get('grade') or data.get('target_grade') or 3)
    out = tamil_lesson_plan(words, concepts, target_grade=grade, model=cfg['model'], base_url=cfg['base_url'])
    return jsonify({'ok': True, 'mode': 'local_ollama', 'model': cfg['model'], 'result': out})


@app.route('/api/ai/questions', methods=['POST'])
def api_ai_questions():
    cfg, err = _require_ai()
    if err: return err
    data = request.get_json(silent=True) or {}
    text = (data.get('text') or '').strip()
    grade = int(data.get('grade') or data.get('target_grade') or 3)
    count = int(data.get('count') or 5)
    if not text:
        return jsonify({'error': 'No text provided'}), 400
    out = tamil_questions(text[:6000], target_grade=grade, count=count, model=cfg['model'], base_url=cfg['base_url'])
    return jsonify({'ok': True, 'mode': 'local_ollama', 'model': cfg['model'], 'result': out})



def _safe_json_loads(value, default):
    try:
        if not value:
            return default
        return json.loads(value)
    except Exception:
        return default


def _compact_analysis_for_ai(row, mode='author'):
    """Build a small, safe prompt payload from a saved analysis.

    We intentionally do NOT send the whole book/PDF to the local model.
    Only the top difficult words, concepts, sentence examples, and offline
    scores are sent, so Analyze remains fast and local Ollama does not get
    overloaded on long books.
    """
    results = _safe_json_loads(row['results_json'] if 'results_json' in row.keys() else None, [])
    meaning = _safe_json_loads(row['meaning_json'] if 'meaning_json' in row.keys() else None, {})
    suitability = _safe_json_loads(row['suitability_json'] if 'suitability_json' in row.keys() else None, {})
    v27 = _safe_json_loads(row['v27_json'] if 'v27_json' in row.keys() else None, {})

    try:
        target_grade = int(v27.get('recommended_grade') or suitability.get('recommended_grade') or 3)
    except Exception:
        target_grade = 3

    # Pick the grade row closest to recommended grade.
    grade_row = next((r for r in results if int(r.get('grade', -1)) == target_grade), None) or (results[-1] if results else {})
    difficult_words = (grade_row.get('unknown_word_list') or grade_row.get('new_word_list') or [])[:25]

    flagged = meaning.get('flagged') or []
    concepts = []
    for f in flagged[:25]:
        item = f.get('item') or f.get('word') or ''
        concept = f.get('concept') or 'general'
        level = f.get('level') or ''
        if item:
            concepts.append({'item': item, 'concept': concept, 'level': level})

    rewrite_items = v27.get('rewrite_suggestions') or []
    long_sentences = []
    for r in rewrite_items:
        if r.get('type') == 'split_sentence' and r.get('original'):
            long_sentences.append(r.get('original'))
        if len(long_sentences) >= 5:
            break

    clusters = []
    for c in (v27.get('concept_clusters') or [])[:6]:
        words = []
        for w in c.get('words', [])[:8]:
            words.append(w.get('word') if isinstance(w, dict) else str(w))
        clusters.append({'label': c.get('label') or c.get('concept'), 'words': [x for x in words if x]})

    glossary = []
    for g in (v27.get('smart_glossary') or [])[:20]:
        glossary.append({'word': g.get('word'), 'meaning': g.get('meaning'), 'level': g.get('class_level')})

    return {
        'book_name': row['book_name'] if 'book_name' in row.keys() else 'Book',
        'target_grade': target_grade,
        'difficulty_score': v27.get('difficulty_score_10'),
        'difficulty_label': v27.get('difficulty_label'),
        'support_level': v27.get('support_level'),
        'independent_reading': v27.get('independent_reading'),
        'known_pct': grade_row.get('comprehension_pct') or grade_row.get('known_pct'),
        'difficult_words': difficult_words,
        'concepts': concepts,
        'concept_clusters': clusters,
        'long_sentences': long_sentences,
        'glossary': glossary,
        'mode': mode,
    }


def _format_ai_payload_for_prompt(payload):
    parts = []
    parts.append(f"Book: {payload.get('book_name')}")
    parts.append(f"Recommended/target class: Std {payload.get('target_grade')}")
    parts.append(f"Difficulty: {payload.get('difficulty_score')}/10 ({payload.get('difficulty_label')})")
    parts.append(f"Teacher support level: {payload.get('support_level')}")
    parts.append(f"Known vocabulary: {payload.get('known_pct')}%")
    if payload.get('difficult_words'):
        parts.append('Difficult words: ' + ', '.join(map(str, payload['difficult_words'][:25])))
    if payload.get('concepts'):
        parts.append('Advanced concepts: ' + '; '.join([f"{c.get('item')} ({c.get('concept')}, Std {c.get('level')})" for c in payload['concepts'][:15]]))
    if payload.get('concept_clusters'):
        parts.append('Concept clusters: ' + '; '.join([f"{c.get('label')}: {', '.join(c.get('words') or [])}" for c in payload['concept_clusters'][:6]]))
    if payload.get('long_sentences'):
        parts.append('Long sentence examples: ' + ' | '.join(payload['long_sentences'][:4]))
    if payload.get('glossary'):
        parts.append('Glossary candidates: ' + '; '.join([f"{g.get('word')}: {g.get('meaning')}" for g in payload['glossary'][:12]]))
    return '\n'.join(parts)


def _run_analysis_ai_enrichment(payload, cfg, mode='author'):
    summary = _format_ai_payload_for_prompt(payload)
    grade = payload.get('target_grade') or 3
    if mode == 'teacher':
        prompt = f"""நீங்கள் தமிழ் ஆசிரியர். கீழே உள்ள புத்தக வாசிப்பு-பகுப்பாய்வு சுருக்கத்தைப் பார்த்து {grade}ஆம் வகுப்பு மாணவர்களுக்கு உதவும் நடைமுறை ஆசிரியர் வழிகாட்டியை உருவாக்குங்கள்.

விதிகள்:
- தமிழில் மட்டும் பதிலளிக்கவும்.
- முழு புத்தகம் கொடுக்கப்படவில்லை; கீழே உள்ள பகுப்பாய்வு சுருக்கத்தின் அடிப்படையில் மட்டும் பதிலளிக்கவும்.
- குறுகிய, பயன்படும் புள்ளிகளாக எழுதவும்.
- பிரிவுகள்: 1) பயன்படுத்தலாமா? 2) வாசிப்புக்கு முன் 3) வாசிக்கும் போது 4) வாசித்த பின் 5) 5 கேள்விகள்.

பகுப்பாய்வு சுருக்கம்:
{summary}

ஆசிரியர் வழிகாட்டி:"""
    else:
        prompt = f"""நீங்கள் தமிழ் குழந்தைகள் புத்தக ஆசிரியருக்கு உதவும் தொகுப்பாசிரியர். கீழே உள்ள புத்தக வாசிப்பு-பகுப்பாய்வு சுருக்கத்தை வைத்து ஆசிரியருக்கு/எழுத்தாளருக்கு மேம்பாட்டு பரிந்துரைகள் தருங்கள்.

விதிகள்:
- தமிழில் மட்டும் பதிலளிக்கவும்.
- முழு புத்தகத்தை மறுஎழுத வேண்டாம்.
- கடின சொற்கள், நீளமான வாக்கியங்கள், கருத்து விளக்கம், glossary ஆகியவற்றுக்கு செயல் படிகள் கொடுக்கவும்.
- பிரிவுகள்: 1) மொத்த மதிப்பீடு 2) எளிமைப்படுத்த வேண்டியவை 3) Glossary சேர்க்க வேண்டியவை 4) மாதிரி எளிய விளக்கங்கள் 5) அடுத்த திருத்தப் படிகள்.

பகுப்பாய்வு சுருக்கம்:
{summary}

ஆசிரியர்/எழுத்தாளர் பரிந்துரைகள்:"""
    return generate(prompt, model=cfg['model'], base_url=cfg['base_url'], temperature=0.2, timeout=240)


@app.route('/api/ai/enrich_analysis', methods=['POST'])
def api_ai_enrich_analysis():
    """Optional local AI enrichment for Analyze Book results.

    Runs only after the normal offline analysis is complete and only when the
    admin/user enabled Ollama in Settings. The endpoint uses a compact summary,
    never the full book text, so it is practical on local 7B models.
    """
    cfg, err = _require_ai()
    if err:
        return err
    data = request.get_json(silent=True) or {}
    analysis_id = data.get('analysis_id')
    mode = (data.get('mode') or 'author').strip().lower()
    if mode not in ('author', 'teacher'):
        mode = 'author'
    if not analysis_id:
        return jsonify({'error': 'analysis_id is required'}), 400

    cache_dir = os.path.join('data', 'cache', 'ai')
    os.makedirs(cache_dir, exist_ok=True)
    cache_path = os.path.join(cache_dir, f'analysis_{analysis_id}_{mode}_{cfg["model"].replace(":", "_").replace("/", "_")}.json')
    if os.path.exists(cache_path):
        try:
            with open(cache_path, 'r', encoding='utf-8') as f:
                cached = json.load(f)
            cached['cached'] = True
            return jsonify(cached)
        except Exception:
            pass

    conn = get_db()
    row = conn.execute('SELECT * FROM analyses WHERE id=?', (analysis_id,)).fetchone()
    conn.close()
    if not row:
        return jsonify({'error': 'Analysis not found'}), 404

    payload = _compact_analysis_for_ai(row, mode=mode)
    try:
        result = _run_analysis_ai_enrichment(payload, cfg, mode=mode)
        out = {'ok': True, 'mode': mode, 'model': cfg['model'], 'result': result, 'payload_summary': payload, 'cached': False}
        with open(cache_path, 'w', encoding='utf-8') as f:
            json.dump(out, f, ensure_ascii=False, indent=2)
        return jsonify(out)
    except Exception as e:
        return jsonify({'error': 'AI enrichment failed', 'details': str(e)}), 500

@app.route('/api/startup_check')
def startup_check():
    """
    Called by the UI on page load.
    Behaviour:
      - Default folder is textbooks_imported/ inside the app directory.
        This is the same folder the Textbook Auto Importer uses when downloading
        from internet sources — so manually placed PDFs and downloaded PDFs
        are all in one place.
      - If config has a different watch_folder set, that takes priority.
      - Only shows the dialog when the user clicks the Rescan button.
        It does not interrupt normal page load, even on an empty database.
      - File counting is 100% local — no internet calls.
    """
    # Default folder: textbooks_imported/ inside the app directory
    # (same folder the Textbook Auto Importer tab downloads into)
    app_dir     = os.path.dirname(os.path.abspath(__file__))
    default_dir = os.path.join(app_dir, 'textbooks_imported')

    cfg    = _fw.load_config()
    folder = cfg.get('watch_folder', '').strip() or default_dir

    # If watch_folder was empty, persist the resolved default so Settings
    # shows the correct path and the watcher starts on next launch
    if not cfg.get('watch_folder', '').strip():
        cfg['watch_folder'] = default_dir
        _fw.save_config(cfg)

    # Check if DB has ever been loaded (first startup detection)
    conn = get_db()
    loaded_count = conn.execute('SELECT COUNT(*) FROM grade_files').fetchone()[0]
    conn.close()
    is_first_startup = (loaded_count == 0)

    # Check rescan flag (set by the Rescan button in the UI)
    rescan_requested = cfg.get('rescan_requested', False)
    if rescan_requested:
        cfg['rescan_requested'] = False
        _fw.save_config(cfg)

    scan_active = _fw.WATCHER_STATUS.get('scan_active', False)

    # Only proceed with file counting if the user explicitly asked.
    should_show = rescan_requested and not scan_active

    if not should_show:
        return jsonify({
            'folder':           folder,
            'show_dialog':      False,
            'is_first_startup': is_first_startup,
            'new_files':        0,
            'total_files':      0,
            'scan_active':      scan_active,
        })

    # Count new files — walk subdirectories (files may be in Class_01/, Class_02/ etc.)
    folder_exists = os.path.isdir(folder)
    new_files = []
    all_count = 0

    if folder_exists:
        loaded_db = _get_loaded_hashes()
        with _hash_cache_lock:
            loaded_cache = dict(_hash_cache)
        known_hashes = set(loaded_db.values()) | set(loaded_cache.values())

        exts = {'.pdf', '.txt', '.docx'}
        try:
            for root, dirs, files in os.walk(folder):
                # Mirror scan_folder: don't go deeper than 6 levels
                depth = len(os.path.relpath(root, folder).split(os.sep))
                if depth > 6:
                    dirs.clear()
                    continue
                # Skip hidden dirs
                dirs[:] = [d for d in sorted(dirs) if not d.startswith('.')]
                for fname in sorted(files):
                    ext = os.path.splitext(fname)[1].lower()
                    if ext not in exts:
                        continue
                    fp = os.path.join(root, fname)
                    try:
                        size = os.path.getsize(fp)
                        if size < 10 * 1024:   # skip tiny/temp files
                            continue
                        all_count += 1
                        h = _fw.file_hash(fp)
                        if h not in known_hashes:
                            # Show relative path so user knows which subfolder
                            rel = os.path.relpath(fp, folder)
                            new_files.append({
                                'filename': rel,
                                'size_mb':  round(size / 1024 / 1024, 1),
                            })
                    except OSError:
                        pass
        except OSError:
            pass

    total_mb = sum(f['size_mb'] for f in new_files)
    est_mins = max(1, round(total_mb * 1.2))

    return jsonify({
        'folder':           folder,
        'folder_exists':    folder_exists,
        'show_dialog':      True,
        'is_first_startup': is_first_startup,
        'new_files':        len(new_files),
        'total_files':      all_count,
        'total_mb':         round(total_mb, 1),
        'est_mins':         est_mins,
        'files':            new_files[:20],
        'scan_active':      scan_active,
    })


@app.route('/api/request_rescan', methods=['POST'])
def request_rescan():
    """
    Called by the Rescan button. Sets a flag so the next startup_check
    call returns show_dialog=True with the current file count.
    """
    cfg = _fw.load_config()
    cfg['rescan_requested'] = True
    _fw.save_config(cfg)
    return jsonify({'ok': True})

if __name__ == '__main__':
    # Clear compiled .pyc cache so updates take effect immediately
    import shutil
    _pyc = os.path.join(os.path.dirname(os.path.abspath(__file__)), '__pycache__')
    if os.path.exists(_pyc):
        shutil.rmtree(_pyc)

    os.makedirs('uploads', exist_ok=True)
    init_db()

    import sys
    if '--build-meaning' in sys.argv or '--update-meaning' in sys.argv:
        print('  Building meaning-level knowledge base from existing textbook database...')
        try:
            meta = _meaning_kb.build_from_existing_db(
                DB_PATH, 'data/meaning_kb',
                extract_text_fn=extract_text,
                tokenize_fn=tokenize_tamil,
                stem_fn=get_stem,
                full_rebuild=True,
            )
            print('  ✓ Meaning KB built')
            print(f"  Words: {meta.get('word_count', 0):,} | Phrases: {meta.get('phrase_count', 0):,} | Concepts: {meta.get('concept_count', 0):,}")
            print(f"  Files used: {meta.get('source_files_used_count', 0)} | Missing originals: {meta.get('source_files_missing_count', 0)}")
            sys.exit(0)
        except Exception as e:
            print('  ✗ Meaning KB build failed:', e)
            sys.exit(1)

    # Start folder watcher if configured
    cfg = _fw.load_config()
    folder = cfg.get('watch_folder', '').strip()
    if folder and os.path.isdir(folder):
        WATCH_FOLDER = folder
        print(f'  Watch folder: {folder}')
        t = threading.Thread(target=_fw.scan_folder,
                             args=(folder, cfg, _watcher_process_fn, _get_loaded_hashes),
                             daemon=True)
        t.start()
        _fw.start_watcher(folder, cfg, _watcher_process_fn, _get_loaded_hashes)
    else:
        print('  Watch folder: not configured (set via Settings tab)')

    print('\n  Tamil Book Readability Analyzer (v28)')
    print('  =======================================')
    print(f'  Snowball stemmer: {"✓" if SNOWBALL_AVAILABLE else "✗ — pip install snowballstemmer"}')
    print('  Open: http://localhost:5000\n')
    app.run(debug=False, port=5000)
